import hashlib
import json
import logging
import os
import pathlib
import random
import re
import sys
import time
import unicodedata
import uuid

import google.generativeai as genai
import markdown
import matplotlib
import matplotlib.pyplot as plt
import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt
from dotenv import load_dotenv
from lxml import html
import requests  # Added for Ollama
from transformers import AutoTokenizer  # Added for Ollama client-side tokenization
from PIL import Image
from random_words import RandomWords

# Configure Matplotlib to use LaTeX and load amsmath
matplotlib.rcParams["text.usetex"] = True
matplotlib.rcParams[
    "text.latex.preamble"
] = r"\usepackage{amsmath}  \usepackage{amssymb}"

# Ensure matplotlib doesn't try to use a GUI backend
plt.switch_backend("Agg")

# --- Constants ---
# Multiplier for inline math height based on font size
DEFAULT_INLINE_MATH_HEIGHT_MULTIPLIER = 1.05
# Height in inches for display math images
DEFAULT_DISPLAY_MATH_HEIGHT_INCHES = 0.375
# Default tone if not specified
DEFAULT_WRITING_TONE = "academic, informative, yet engaging"

# Configure logging
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)


# --- Configuration and Environment Loading ---
def load_config(config_path="config.yaml"):
    """Loads the configuration file."""
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f)
        logging.info(f"Configuration loaded successfully from {config_path}")
        return config
    except FileNotFoundError:
        logging.error(f"Error: Configuration file not found at {config_path}")
        sys.exit(1)
    except yaml.YAMLError as e:
        logging.error(f"Error parsing configuration file {config_path}: {e}")
        sys.exit(1)


def render_latex_to_image(
    latex_string_with_delimiters, equation_image_dir, is_display_style=False
):
    """
    Renders a LaTeX math string (potentially including delimiters like \(...\) or $$...$$)
    to a PNG image using matplotlib's mathtext.
    Adjusts padding based on whether it's display or inline style.
    """
    # 1. Clean the input string: Remove common LaTeX math delimiters robustly
    raw_latex_code = latex_string_with_delimiters.strip()
    original_input_for_log = raw_latex_code  # Keep original for logging

    # --- Improved Delimiter Removal ---
    # Repeatedly strip common delimiters until none are found at the ends
    cleaned = True
    while cleaned:
        cleaned = False
        original_length = len(raw_latex_code)
        if raw_latex_code.startswith(r"\(") and raw_latex_code.endswith(r"\)"):
            raw_latex_code = raw_latex_code[2:-2].strip()
            logging.debug("Stripped '\\(...\\)' delimiters.")
        elif raw_latex_code.startswith(r"\[") and raw_latex_code.endswith(r"\]"):
            raw_latex_code = raw_latex_code[2:-2].strip()
            logging.debug("Stripped '\\[...\\]' delimiters.")
        elif raw_latex_code.startswith("$$") and raw_latex_code.endswith("$$"):
            raw_latex_code = raw_latex_code[2:-2].strip()
            logging.debug("Stripped '$$...$$' delimiters.")
        elif (
            raw_latex_code.startswith("$")
            and raw_latex_code.endswith("$")
            and not raw_latex_code.startswith(
                "$$"
            )  # Avoid stripping single $ from $$..$$
        ):
            raw_latex_code = raw_latex_code[1:-1].strip()
            logging.debug("Stripped '$...$' delimiters.")

        if len(raw_latex_code) < original_length:
            cleaned = True  # Mark that we changed something and should loop again

    if not raw_latex_code:
        logging.warning(
            f"LaTeX string became empty after removing delimiters. Original: '{original_input_for_log}'"
        )
        return None
    # --- End Improved Delimiter Removal ---

    # 2. Wrap the *cleaned* LaTeX code with $ for mathtext
    mathtext_string = f"${raw_latex_code}$"
    logging.debug(f"Attempting to render mathtext: {mathtext_string}")

    # Generate a unique filename
    filename = f"eq_{uuid.uuid4().hex}.png"
    # Use the passed-in directory
    filepath = equation_image_dir / filename

    fig = plt.figure()
    try:
        # Render the LaTeX using mathtext
        fontsize = 12 if not is_display_style else 14
        fig.text(
            0, 0, mathtext_string, fontsize=fontsize, va="bottom", math_fontfamily="cm"
        )

        # Adjust figure size to tightly fit the text
        renderer = fig.canvas.get_renderer()
        try:
            bbox = fig.texts[0].get_window_extent(renderer=renderer)
            if bbox.width <= 1 or bbox.height <= 1:
                raise ValueError(
                    f"Rendered bounding box is too small ({bbox.width}x{bbox.height}), likely mathtext render error."
                )
        except IndexError:
            raise ValueError(
                "Matplotlib fig.text failed to render (no text objects found)."
            )
        except ValueError as ve:
            raise ValueError(f"Matplotlib rendering error: {ve}")

        dpi = fig.get_dpi()
        # Calculate base width/height from bbox, add a minimal buffer to prevent clipping
        base_width = (bbox.width / dpi) + 0.01
        base_height = (bbox.height / dpi) + 0.01

        # --- Determine padding based on style ---
        if is_display_style:
            # Keep existing padding for display math
            figure_padding = 0.1
            save_pad_inches = 0.05
            logging.debug("Using display style padding.")
        else:
            # Use zero padding for inline math
            figure_padding = 0.0
            save_pad_inches = 0.0
            logging.debug("Using zero padding for inline style.")
        # --- End padding determination ---

        # Set final figure size including the calculated padding
        final_width = base_width + figure_padding
        final_height = base_height + figure_padding
        fig.set_size_inches(final_width, final_height)

        # Calculate text position within the final figure size
        # If figure_padding is 0, text starts near bottom-left (considering the small base buffer)
        text_x = (figure_padding / 2) / final_width if final_width > 0 else 0
        text_y = (figure_padding / 2) / final_height if final_height > 0 else 0
        fig.texts[0].set_position((text_x, text_y))

        # Make sure the directory exists (it should be created in main, but double-check)
        equation_image_dir.mkdir(parents=True, exist_ok=True)

        # Save the figure using the determined save padding
        plt.savefig(
            filepath,
            dpi=300,
            bbox_inches="tight",  # Still use tight bbox to help
            pad_inches=save_pad_inches,  # Use conditional padding here
            transparent=True,
        )
        plt.close(fig)
        logging.debug(f"Rendered LaTeX to {filepath}")
        return str(filepath)
    except Exception as e:
        # Log the original input (with delimiters) and the final mathtext string
        logging.error(
            f"Failed to render LaTeX input: '{original_input_for_log}'. Processed mathtext: '{mathtext_string}'. Error: {e}"
        )
        plt.close(fig)
        return None


# --- Utility Functions ---
def sanitize_filename(filename_base, max_length=200):
    """Converts a string into a safe filename stem (without extension)."""
    filename_base = str(filename_base)  # Ensure it's a string
    filename_base = (
        unicodedata.normalize("NFKD", filename_base)
        .encode("ascii", "ignore")
        .decode("ascii")
    )
    # Replace whitespace and common invalid chars with underscore
    filename_base = re.sub(r"[\s/\\:?*\"<>|]+", "_", filename_base)
    # Remove any remaining non-alphanumeric characters (except underscore, hyphen, period)
    filename_base = re.sub(r"[^\w\-_\.]", "", filename_base)
    # Remove leading/trailing underscores, periods, hyphens
    filename_base = filename_base.strip("._-")
    # Limit length
    filename_base = filename_base[:max_length]
    # Handle empty string after sanitization
    if not filename_base:
        # Use a generic placeholder if sanitization results in empty string
        return "sanitized_empty"
    # Remove trailing periods (can cause issues on Windows)
    filename_base = filename_base.rstrip(".")
    # Ensure it's not empty again after stripping trailing period
    if not filename_base:
        return "sanitized_empty_final"
    return filename_base


def setup_environment():
    """Loads environment variables from .env file and retrieves API key."""
    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        logging.error(
            "Error: GEMINI_API_KEY not found in .env file or environment variables."
        )
        sys.exit(1)
    logging.info("Environment variables loaded and API key found.")
    return api_key


def configure_gemini(api_key):
    """Configures the Google Generative AI client."""
    try:
        genai.configure(api_key=api_key)
        logging.info("Google Generative AI client configured.")
    except Exception as e:
        logging.error(f"Error configuring Google Generative AI: {e}")
        sys.exit(1)


# --- Caching Mechanism ---
def get_cache_path(prompt_text, cache_dir, cache_prefix=None):
    """Generates the cache file path, optionally prepending a prefix."""
    prompt_hash = hashlib.sha256(prompt_text.encode("utf-8")).hexdigest()
    pathlib.Path(cache_dir).mkdir(parents=True, exist_ok=True)

    filename_base = prompt_hash
    if cache_prefix:
        # Sanitize the prefix to make it filename-safe and limit length
        # Use a shorter length limit for prefixes to avoid overly long filenames
        sanitized_prefix = sanitize_filename(cache_prefix, max_length=50)
        if sanitized_prefix:  # Ensure sanitization didn't result in an empty string
            filename_base = f"{sanitized_prefix}_{prompt_hash}"
            logging.debug(f"Using cache prefix: '{sanitized_prefix}'")
        else:
            logging.warning(
                f"Cache prefix '{cache_prefix}' sanitized to empty string. Using hash only."
            )

    return pathlib.Path(cache_dir) / f"{filename_base}.json"


def load_from_cache(prompt_text, cache_dir, cache_prefix=None):
    """Loads response from cache if available."""
    cache_file = get_cache_path(prompt_text, cache_dir, cache_prefix)
    if cache_file.exists():
        try:
            with open(cache_file, "r", encoding="utf-8") as f:
                cached_data = json.load(f)
            if "prompt" in cached_data and "response" in cached_data:
                # Log the actual filename for clarity
                logging.info(f"Cache hit for file: {cache_file.name}")
                return cached_data["response"]
            else:
                logging.warning(f"Invalid cache file format: {cache_file}. Ignoring.")
                return None
        except Exception as e:
            logging.error(f"Error reading cache file {cache_file}: {e}")
            return None
    logging.debug(f"Cache miss for file: {cache_file.name}")
    return None


def save_to_cache(prompt_text, response_text, cache_dir, cache_prefix=None):
    """Saves the API response to the cache."""
    cache_file = get_cache_path(prompt_text, cache_dir, cache_prefix)
    try:
        cache_data = {"prompt": prompt_text, "response": response_text}
        with open(cache_file, "w", encoding="utf-8") as f:
            json.dump(cache_data, f, ensure_ascii=False, indent=4)
        logging.info(f"Response saved to cache: {cache_file}")
    except Exception as e:
        logging.error(f"Error saving response to cache file {cache_file}: {e}")


# --- LLM API Interaction ---
def _call_gemini_api_internal(prompt, config, cache_prefix=None):
    """
    Internal function to call the Gemini API.
    Assumes caching is handled by the caller.

    Args:
        prompt (str): The prompt to send to the API.
        config (dict): The application configuration.
        cache_prefix (str, optional): A prefix to add to the cache filename
                                      for better identification. Defaults to None.
                                      (Note: cache_prefix is for logging/context here, actual caching is external)
    Returns:
        str or None: The API response text, or None if an error occurred.
    """
    api_settings_conf = config.get("api_settings", {})
    gemini_conf = api_settings_conf.get("gemini", {})

    default_max_retries = api_settings_conf.get("default_max_retries", 3)
    default_retry_delay = api_settings_conf.get("default_retry_delay_seconds", 5)

    # Logging for Gemini-specific call initiation (cache prefix is for context)
    model_name = gemini_conf.get("model", "gemini-2.0-flash-latest")
    temperature = float(
        gemini_conf.get("temperature", 1.0)
    )  # Default from example config.yaml

    max_retries = int(gemini_conf.get("max_retries", default_max_retries))
    retry_delay = int(gemini_conf.get("retry_delay_seconds", default_retry_delay))
    # safety_settings would be fetched from gemini_conf if specified in config.yaml under api_settings.gemini
    safety_settings = gemini_conf.get("safety_settings", None)

    verbose_debug = config.get("debug_options", {}).get("verbose_debug", False)
    stream_gemini = (
        verbose_debug  # Specifically for Gemini streaming if verbose_debug is on
    )

    try:
        if verbose_debug:
            logging.info(f"Gemini API Prompt for model '{model_name}':\n{prompt}")
            # For very long prompts, you might want to log only a portion or a summary
            # logging.info(f"Gemini API Prompt for model '{model_name}' (first 500 chars):\n{prompt[:500]}...")

        model = genai.GenerativeModel(model_name)
        generation_config = genai.types.GenerationConfig(temperature=temperature)

        # Count tokens for Gemini prompt
        try:
            token_count_response = model.count_tokens(prompt)
            prompt_token_count = token_count_response.total_tokens
            logging.info(
                f"Gemini prompt token count for model '{model_name}': {prompt_token_count} tokens."
            )
        except Exception as e_token:
            logging.warning(
                f"Could not count tokens for Gemini prompt (model '{model_name}'): {e_token}"
            )

        for attempt in range(max_retries):
            try:
                response = model.generate_content(
                    prompt,
                    generation_config=generation_config,
                    safety_settings=safety_settings,
                    stream=stream_gemini,
                )

                if stream_gemini:
                    logging.info(f"Streaming Gemini response for model '{model_name}':")
                    full_response_text_parts = []
                    print(f"\n--- Gemini Stream ({model_name}) ---")
                    for chunk in response:
                        if chunk.parts:
                            response_part = chunk.text
                            print(
                                response_part, end="", flush=True
                            )  # Stream to console
                            full_response_text_parts.append(response_part)
                        elif (
                            chunk.prompt_feedback and chunk.prompt_feedback.block_reason
                        ):  # Prompt itself is blocked
                            logging.warning(
                                f"Gemini API stream blocked. Reason: {chunk.prompt_feedback.block_reason}"
                            )
                            print(f"\n--- End Gemini Stream (Blocked) ---")
                            return None  # Blocked, don't retry
                        # We don't typically get 'done' in the same way as Ollama,
                        # the stream just ends. The loop finishing means it's done.

                    print(f"\n--- End Gemini Stream (Done) ---")
                    logging.info(
                        f"Gemini API stream completed for model '{model_name}'."
                    )
                    final_text = "".join(full_response_text_parts).strip()
                    return final_text
                else:  # Not streaming
                    response_text = None  # Initialize

                    # Handle prompt feedback first: if the prompt itself was blocked, don't retry.
                    if (
                        response.prompt_feedback
                        and response.prompt_feedback.block_reason
                    ):
                        logging.warning(
                            f"Gemini API call blocked due to prompt. Reason: {response.prompt_feedback.block_reason}. Will not retry."
                        )
                        return None  # Explicitly do not retry prompt blocks

                    # If prompt was not blocked, check for response parts
                    if response.parts:
                        response_text = response.text
                    # Removed the elif for prompt_feedback.block_reason here as it's handled above.
                    # The case below is for when there are no parts, and it wasn't a prompt block.
                    # This could be due to finish_reason (e.g., SAFETY on response, MAX_TOKENS).
                    else:
                        # Check finish reason even if parts are empty
                        finish_reason = "UNKNOWN"
                        try:
                            # Access finish_reason safely
                            if response.candidates:
                                finish_reason = response.candidates[
                                    0
                                ].finish_reason.name  # Use .name for enum
                        except (AttributeError, IndexError):
                            logging.warning(
                                "Could not determine finish reason from response."
                            )

                        logging.warning(
                            f"API call returned no content or parts (and was not prompt-blocked). Finish Reason: {finish_reason}. Will attempt retry if applicable."
                        )
                        # response_text remains None

                    if response_text is not None:  # Check if we got valid text
                        logging.info(
                            f"Gemini API call successful for model {model_name}."
                        )
                        return response_text
                    else:
                        logging.warning(
                            f"API attempt {attempt + 1} for model {model_name} resulted in no content (response_text is None). Will proceed to retry logic."
                        )

            except genai.types.BlockedPromptException as bpe:
                logging.error(
                    f"Gemini API call attempt {attempt + 1} for model {model_name} failed due to a blocked prompt: {bpe}. Will not retry."
                )
                return None  # Do not retry if the prompt itself is blocked
            except genai.types.StopCandidateException as sce:
                # This exception means the response generation was stopped (e.g., safety, recitation).
                # Retrying might yield a different result, especially with temperature > 0.
                logging.warning(
                    f"Gemini API call attempt {attempt + 1} for model {model_name} stopped during candidate generation: {sce}. Retrying..."
                )
                # No 'return None' here, so it falls through to the retry delay logic.
            except Exception as e:
                logging.warning(
                    f"Gemini API call attempt {attempt + 1} for model {model_name} failed: {e}"
                )
                if "quota" in str(e).lower():  # Basic check for quota issues
                    logging.warning(
                        f"Gemini API quota likely exceeded: {e}. Retrying as per configuration..."
                    )
                    # Removed 'return None' to allow retry for quota issues
                # For other general exceptions, the loop will continue to the retry logic.

            if attempt < max_retries - 1:
                logging.info(f"Retrying Gemini API call in {retry_delay} seconds...")
                time.sleep(retry_delay)

        logging.error(
            f"Gemini API call for model {model_name} failed after {max_retries} attempts."
        )
        return None  # Explicitly return None after all retries fail

    except Exception as e:
        logging.error(f"An unexpected error occurred during Gemini API call setup: {e}")
        return None


def _call_ollama_api_internal(prompt, config, cache_prefix=None):
    """
    Internal function to call the Ollama API.
    Assumes caching is handled by the caller.

    Args:
        prompt (str): The prompt to send to the API.
        config (dict): The application configuration.
        cache_prefix (str, optional): Contextual prefix, caching is external.

    Returns:
        str or None: The API response text, or None if an error occurred.
    """
    api_settings_conf = config.get("api_settings", {})
    ollama_config = api_settings_conf.get("ollama", {})

    default_max_retries = api_settings_conf.get("default_max_retries", 3)
    default_retry_delay = api_settings_conf.get("default_retry_delay_seconds", 5)

    base_url = ollama_config.get("base_url", "http://localhost:11434")
    model_name = ollama_config.get("model", "llama3")  # Default Ollama model
    temperature = float(ollama_config.get("temperature", 0.7))
    tokenizer_model_name = ollama_config.get(
        "tokenizer_model", "NousResearch/Llama-3-8B-Instruct-hf"
    )  # Default Llama3 tokenizer

    max_retries = int(ollama_config.get("max_retries", default_max_retries))
    retry_delay = int(ollama_config.get("retry_delay_seconds", default_retry_delay))
    request_timeout = int(
        ollama_config.get("request_timeout_seconds", 120)
    )  # Default 2 mins
    api_url = f"{base_url.rstrip('/')}/api/generate"

    verbose_debug = config.get("debug_options", {}).get("verbose_debug", False)
    stream_ollama = (
        verbose_debug  # Specifically for Ollama streaming if verbose_debug is on
    )

    payload = {
        "model": model_name,
        "prompt": prompt,
        "stream": stream_ollama,
        "options": {"temperature": temperature},
    }

    if verbose_debug:
        logging.info(f"Ollama API Prompt for model '{model_name}':\n{prompt}")
        # For very long prompts, you might want to log only a portion or a summary
        # logging.info(f"Ollama API Prompt for model '{model_name}' (first 500 chars):\n{prompt[:500]}...")

    # Attempt client-side token counting for Ollama
    # Note: For higher efficiency with many calls, consider loading the tokenizer once outside this function.
    if tokenizer_model_name:
        try:
            logging.debug(
                f"Loading tokenizer: {tokenizer_model_name} for Ollama prompt token count."
            )
            tokenizer = AutoTokenizer.from_pretrained(tokenizer_model_name)
            token_ids = tokenizer.encode(prompt)
            num_tokens = len(token_ids)
            logging.info(
                f"Ollama client-side token count for prompt (tokenizer: '{tokenizer_model_name}', model: '{model_name}'): {num_tokens} tokens."
            )
        except Exception as e_token_ollama:
            logging.warning(
                f"Could not count tokens for Ollama prompt using tokenizer '{tokenizer_model_name}' (model: '{model_name}'): {e_token_ollama}"
            )
            logging.info(
                f"Ollama API for model '{model_name}': Standard Ollama API does not provide a direct prompt token count. Client-side estimation failed."
            )
    else:
        logging.info(
            f"Ollama API for model '{model_name}': No tokenizer_model configured for client-side token counting. Standard Ollama API does not provide a direct prompt token count."
        )

    for attempt in range(max_retries):
        try:
            response = requests.post(
                api_url,
                headers={"Content-Type": "application/json"},
                json=payload,
                timeout=request_timeout,
                stream=stream_ollama,  # Pass stream=True to requests.post if streaming
            )
            response.raise_for_status()  # Raises HTTPError for bad responses (4XX, 5XX)

            if stream_ollama:
                logging.info(f"Streaming Ollama response for model '{model_name}':")
                full_response_text_parts = []
                print(f"\n--- Ollama Stream ({model_name}) ---")
                for line in response.iter_lines():
                    if line:
                        decoded_line = line.decode("utf-8")
                        try:
                            chunk = json.loads(decoded_line)
                            if "error" in chunk:
                                logging.error(
                                    f"Ollama API error during stream for model '{model_name}': {chunk['error']}"
                                )
                                print(f"\n--- End Ollama Stream (Error) ---")
                                return None  # Specific Ollama error, don't retry

                            response_part = chunk.get("response", "")
                            print(
                                response_part, end="", flush=True
                            )  # Stream to console
                            full_response_text_parts.append(response_part)

                            if chunk.get("done"):
                                print(f"\n--- End Ollama Stream (Done) ---")
                                logging.info(
                                    f"Ollama API stream completed for model '{model_name}'."
                                )
                                final_text = "".join(full_response_text_parts).strip()
                                return final_text
                        except json.JSONDecodeError:
                            logging.error(
                                f"Error decoding JSON chunk from Ollama stream: {decoded_line}"
                            )
                            print(f"\n--- End Ollama Stream (JSON Error) ---")
                            return None
                # This part might be reached if the stream ends unexpectedly without a 'done: true'
                print(f"\n--- End Ollama Stream (Unexpected End) ---")
                logging.warning("Ollama stream ended without a 'done: true' message.")
                return (
                    "".join(full_response_text_parts).strip()
                    if full_response_text_parts
                    else None
                )
            else:  # Not streaming
                response_data = response.json()
                if "error" in response_data:
                    logging.error(
                        f"Ollama API error for model '{model_name}': {response_data['error']}"
                    )
                    return None
                if "response" in response_data:
                    response_text = response_data["response"]
                    logging.info(
                        f"Ollama API call successful for model '{model_name}'."
                    )
                    return response_text.strip()
                else:
                    logging.warning(
                        f"Ollama API response for model '{model_name}' did not contain 'response' key. Attempt {attempt + 1}/{max_retries}. Data: {response_data}"
                    )

        except requests.exceptions.HTTPError as e:
            logging.warning(
                f"Ollama API call (model '{model_name}') attempt {attempt + 1} failed with HTTPError: {e}. Status: {e.response.status_code}"
            )
            if e.response.status_code == 404:  # Model not found
                try:
                    error_detail = e.response.json().get("error", "Model not found")
                    logging.error(
                        f"Ollama model '{model_name}' not found: {error_detail}. Please ensure the model is pulled and available."
                    )
                except json.JSONDecodeError:
                    logging.error(f"Ollama model '{model_name}' not found (404).")
                return None  # Don't retry if model not found
        except (
            requests.exceptions.RequestException
        ) as e:  # Covers ConnectionError, Timeout, etc.
            logging.warning(
                f"Ollama API call (model '{model_name}') attempt {attempt + 1} failed: {e}"
            )

        if attempt < max_retries - 1:
            logging.info(f"Retrying Ollama API call in {retry_delay} seconds...")
            time.sleep(retry_delay)
        else:
            logging.error(
                f"Ollama API call for model '{model_name}' failed after {max_retries} attempts."
            )
            return None
    return None  # Should be covered by loop logic, but as a safeguard


def call_llm_api(prompt, config, cache_prefix=None):
    """
    Calls the configured LLM API (Gemini or Ollama), using caching.
    """
    cache_dir = config.get(
        "cache_dir", "api_cache"
    )  # This is now topic and model specific
    cached_response = load_from_cache(prompt, cache_dir, cache_prefix)
    if cached_response is not None:
        return cached_response

    api_settings = config.get("api_settings", {})
    api_provider = api_settings.get("provider", "gemini")  # Default to gemini
    logging.info(
        f"Calling {api_provider.upper()} API... (Cache Prefix: {cache_prefix or 'None'})"
    )

    response_text = None
    if api_provider == "gemini":
        response_text = _call_gemini_api_internal(prompt, config, cache_prefix)
    elif api_provider == "ollama":
        response_text = _call_ollama_api_internal(prompt, config, cache_prefix)
    else:
        logging.error(f"Unsupported API provider: {api_provider}")
        return None

    if response_text is not None:
        save_to_cache(prompt, response_text, cache_dir, cache_prefix)
    return response_text


# --- Book Generation Functions ---
def generate_random_gender(config):
    """Randomly returns either "male" or "female"."""
    return random.choice(["male", "female"])


def determine_gender_from_name(config, author_name):
    """Determines the likely gender (male/female) based on the author's name using the Gemini API."""
    logging.info(f"Attempting to determine gender for author name: '{author_name}'...")

    # Check if the name seems valid (basic check)
    if (
        not author_name
        or not isinstance(author_name, str)
        or " " not in author_name.strip()
    ):
        logging.warning(
            f"Invalid or potentially incomplete name provided ('{author_name}'). Cannot reliably determine gender."
        )
        return None  # Indicate failure

    prompt = f"""
Based *only* on the full name '{author_name}', what is the most likely gender associated with the first name?

Consider common associations in Western cultures, primarily English-speaking contexts, as the book is in British English.

Output *only* one of the following words:
- male
- female
- other (use this only if the name is strongly ambiguous, unisex, or clearly not a typical given name)

Do not add any introductory text, explanations, or quotation marks. Just the single word.
"""

    gender_text = call_llm_api(
        prompt,
        config,
        cache_prefix=f"determine_gender_{sanitize_filename(author_name, 30)}",
    )

    if gender_text:
        cleaned_gender = gender_text.strip().lower()
        valid_genders = ["male", "female", "other"]
        if cleaned_gender in valid_genders:
            logging.info(
                f"Successfully determined likely gender for '{author_name}': '{cleaned_gender}'"
            )
            return cleaned_gender
        else:
            logging.warning(
                f"API returned an unexpected value for gender ('{gender_text}'). Treating as undetermined."
            )
            return None  # Indicate failure or ambiguity
    else:
        logging.error(f"Failed to determine gender for '{author_name}' via API.")
        return None  # Indicate failure


def generate_random_name(config, gender):
    """Generates a random author name based on the specified gender using the Gemini API."""
    logging.info(f"Auto-generating random author name for gender: {gender}...")
    prompt = f"""Generate a single, plausible-sounding full name (first and last name)
for a fictional author. The author's gender is {gender}.
The author is notionally writing a book about
'{config['generation_params']['main_topic']}'. The setting of the book is
described as: {config['generation_params']['setting']}. Key concepts
include: {', '.join(config['generation_params']['key_concepts'])}.
Consider a name that might appear on a book. Output *only* the full name.
Do not add introductory text, explanations, or quotation marks."""

    name_text = call_llm_api(prompt, config, cache_prefix="random_name")

    if name_text:
        cleaned_name = name_text.strip().strip("\"'").strip()
        if cleaned_name and " " in cleaned_name:
            logging.info(f"Successfully auto-generated random name: '{cleaned_name}'")
            return cleaned_name
        else:
            logging.warning(
                f"Generated name ('{name_text}') might be invalid (e.g., missing space). Using it anyway."
            )
            return cleaned_name if cleaned_name else None
    else:
        logging.error("Failed to generate random name via API.")
        return None


def generate_random_topic(config):
    """
    Generates a book topic using the Gemini API.
    Uses a seed from config if provided, otherwise generates a random one.
    """
    logging.info("Determining topic seed...")
    if "generation_params" not in config:
        config["generation_params"] = {}

    # Check if a seed is provided in the config
    provided_seed = config["generation_params"].get("random_topic_seed", "").strip()

    if provided_seed:
        random_seed = provided_seed
        logging.info(f"Using random topic seed from config: '{random_seed}'")
    else:
        logging.info("No seed in config, generating a new random topic seed...")
        # Generate a new random seed if none was provided
        try:
            random_seed = " ".join(RandomWords().random_words(count=15))
            # Store the newly generated seed back into the config dictionary (in memory)
            config["generation_params"]["random_topic_seed"] = random_seed
            logging.info(f"Generated random topic seed: '{random_seed}'")
        except Exception as e:
            logging.error(f"Failed to generate random words for seed: {e}")
            # Fallback seed in case RandomWords fails
            random_seed = f"fallback_seed_{uuid.uuid4().hex[:8]}"
            config["generation_params"]["random_topic_seed"] = random_seed
            logging.warning(f"Using fallback seed: '{random_seed}'")

    prompt = f"""Generate a topic for a book.
Random seed: {random_seed}
Output *only* the topic text itself.
Do not add introductory text, explanations, or quotation marks.
Output in British English."""

    logging.info(f"Generating random topic using seed: '{random_seed}'...")
    topic_text = call_llm_api(prompt, config, cache_prefix="random_topic")

    if topic_text:
        cleaned_topic = topic_text.strip().strip("\"'").rstrip(".").strip()
        if cleaned_topic:
            logging.info(f"Successfully auto-generated random topic: '{cleaned_topic}'")
            return cleaned_topic
        else:
            logging.warning("Generated random topic was empty after cleaning.")
            return None
    else:
        logging.error("Failed to generate random topic via API.")
        return None


def generate_setting(config):
    """Generates the setting using the Gemini API based on the main topic."""
    logging.info("Auto-generating setting...")
    main_topic = config.get("generation_params", {}).get(
        "main_topic", "[No Main Topic Provided]"
    )

    if main_topic == "[No Main Topic Provided]":
        logging.error("Cannot generate setting without 'main_topic' in config.")
        return None

    prompt = f"""Based on the main topic '{main_topic}', generate an
short description of the setting where this topic could be explored.
Output only the setting description text. Do not add introductory text.
Output in British English."""

    setting_text = call_llm_api(prompt, config, cache_prefix="setting")

    if setting_text:
        cleaned_setting = setting_text.strip().strip("\"'")
        if cleaned_setting:
            logging.info(f"Successfully auto-generated setting:\n'{cleaned_setting}'")
            return cleaned_setting
        else:
            logging.warning("Generated setting was empty after cleaning.")
            return None
    else:
        logging.error("Failed to generate setting via API.")
        return None


def generate_writing_tone(config):
    """Generates a suitable writing tone using the Gemini API."""
    logging.info("Auto-generating writing tone...")
    main_topic = config.get("generation_params", {}).get(
        "main_topic", "[No Main Topic Provided]"
    )
    setting = config.get("generation_params", {}).get(
        "setting", "[No  Setting Provided]"
    )
    key_concepts = config.get("generation_params", {}).get("key_concepts", [])

    if main_topic == "[No Main Topic Provided]" or setting == "[No  Setting Provided]":
        logging.warning(
            "Cannot generate specific writing tone without 'main_topic' and 'setting'. Using a generic prompt."
        )
        prompt = f"""Generate some words describing a suitable writing tone
for a book. Output only the phrase describing the tone. Do not add introductory
text. Output in British English."""
    else:
        concepts_str = (
            ", ".join(key_concepts)
            if key_concepts
            else "[No specific concepts provided]"
        )
        prompt = f"""Based on the main topic '{main_topic}', a setting described as:
"{setting}" and key concepts including: {concepts_str}, generate some words
describing the most suitable writing tone for a book exploring this
topic. Output *only* the phrase describing the tone. Do not add introductory
text. Output in British English."""
    tone_text = call_llm_api(prompt, config, cache_prefix="writing_tone")

    if tone_text:
        cleaned_tone = tone_text.strip().strip("\"'").rstrip(".")
        if cleaned_tone:
            logging.info(f"Successfully auto-generated writing tone: '{cleaned_tone}'")
            return cleaned_tone
        else:
            logging.warning("Generated writing tone was empty after cleaning.")
            return None
    else:
        logging.error("Failed to generate writing tone via API.")
        return None


def generate_key_concepts(config):
    """Generates key concepts using the Gemini API based on topic and setting."""
    logging.info("Auto-generating key concepts...")

    main_topic = config.get("generation_params", {}).get(
        "main_topic", "[No Main Topic Provided]"
    )
    setting = config.get("generation_params", {}).get(
        "setting", "[No  Setting Provided]"
    )

    if main_topic == "[No Main Topic Provided]" or setting == "[No  Setting Provided]":
        logging.error(
            "Cannot generate key concepts without 'main_topic' and 'generation_params.setting' in config."
        )
        return None

    prompt = f"""Based on the main topic '{main_topic}' in a setting described as:
"{setting}"

Generate a short list of distinct and relevant key concepts or
terms that would be central to exploring this topic within the setting.

Format the output as a simple comma-separated list. Example:
Concept One, Concept Two, Another Key Term, Fourth Idea, Final Concept

Provide *only* the comma-separated list of concepts. Do not add introductory text.
Output in British English."""

    concepts_text = call_llm_api(prompt, config, cache_prefix="key_concepts")

    if concepts_text:
        cleaned_text = concepts_text.strip().strip("\"'")
        generated_concepts = [
            concept.strip() for concept in cleaned_text.split(",") if concept.strip()
        ]
        return generated_concepts
    else:
        logging.error("Failed to generate key concepts via API.")
        return None


def generate_book_title(config):
    """Generates the book title using the Gemini API. Exits script on failure."""
    logging.info("Generating book title...")
    prompt = f"""Generate the book title for a book about
'{config['generation_params']['main_topic']}'. The setting of the book is
described as: {config['generation_params']['setting']}. Key concepts
include: {', '.join(config['generation_params']['key_concepts'])}.
Do not generate a two-part title. The generated title must not contain a subtitle.
Provide only the title text. Do not add introductory text. The title must not
contain these punctuations: '-' or ':'. Output one title only. Output in
British English."""
    title = call_llm_api(prompt, config, cache_prefix="book_title")
    if title is None:
        logging.error("Fatal: Failed to generate book title after retries. Exiting.")
        sys.exit(1)
    cleaned_title = title.strip().strip("\"'")
    if not cleaned_title:
        logging.error("Fatal: Generated book title is empty. Exiting.")
        sys.exit(1)
    return cleaned_title


def generate_book_subtitle(config, book_title, summary_context):
    """Generates a book subtitle based on the title, topic, and chapter summaries."""
    logging.info(f"Generating subtitle for book: '{book_title}' (using summaries)...")
    gen_params = config.get("generation_params", {})

    # Check if subtitle is provided in config
    subtitle_from_config = gen_params.get("book_subtitle", "").strip()
    if subtitle_from_config:
        logging.info(f"Using book_subtitle from config: '{subtitle_from_config}'")
        return subtitle_from_config

    logging.info(
        "No 'book_subtitle' found in config or it was empty. Attempting to auto-generate one."
    )

    prompt = f"""Generate a subtitle for the book titled '{book_title}'.
The main topic of the book is '{config['generation_params']['main_topic']}'.
The setting of the book described as:
{config['generation_params']['setting']}.
Key concepts include: {', '.join(config['generation_params']['key_concepts'])}.
{summary_context}
The subtitle should complement the main title.
Do not generate a two-part subtitle.
The subtitle must not contain these punctuations: '-' or ':'.
Provide only the subtitle text. Output one subtitle only.
Do not add introductory text. Output in British English."""

    subtitle = call_llm_api(prompt, config, cache_prefix="book_subtitle")
    if subtitle:
        cleaned_subtitle = subtitle.strip().strip("\"'")
        if cleaned_subtitle:
            logging.info(f"Successfully generated subtitle: {cleaned_subtitle}")
            return cleaned_subtitle
        else:
            logging.warning(
                "Generated subtitle was empty after cleaning. No subtitle will be used."
            )
            return None
    else:
        logging.warning("Failed to generate subtitle. No subtitle will be used.")
        return None


def is_book_non_fiction(config, book_title):
    """
    Determines if the book is likely non-fiction using the Gemini API.
    Caches the result in config['generation_params']['is_fiction'].
    Returns True if likely non-fiction, False otherwise (fiction or indeterminate).
    """
    gen_params = config.setdefault(
        "generation_params", {}
    )  # Ensure gen_params exists and can be modified

    # Check if 'is_fiction' is already determined and cached
    if "is_fiction" in gen_params and isinstance(gen_params["is_fiction"], bool):
        is_fiction_cached = gen_params["is_fiction"]
        logging.info(
            f"Using cached 'is_fiction' status for book '{book_title}': {is_fiction_cached}. "
            f"Returning non-fiction status: {not is_fiction_cached}."
        )
        return (
            not is_fiction_cached
        )  # True if non-fiction (is_fiction is False), False if fiction (is_fiction is True)

    # If not cached, proceed with API call
    main_topic = gen_params.get("main_topic", "[No Main Topic Provided]")
    setting = gen_params.get(
        "setting", "[No Setting Provided]"
    )  # Corrected variable name
    key_concepts = gen_params.get("key_concepts", [])

    if main_topic == "[No Main Topic Provided]":
        logging.warning(
            "Cannot determine book type for character list decision: 'main_topic' is missing. "
            "Assuming fiction to be safe and allow character generation if not explicitly disabled."
        )
        # Store this assumption and return
        gen_params["is_fiction"] = True  # Assume fiction
        return False  # is_book_non_fiction returns False for fiction

    concepts_str = (
        ", ".join(key_concepts) if key_concepts else "[No specific concepts provided]"
    )

    prompt = f"""
Based on the following details of a book:
- Title: '{book_title}'
- Main Topic: '{main_topic}'
- Setting: "{setting}"
- Key Concepts: {concepts_str}

Is this book most likely non-fiction?
Answer with only 'yes' or 'no'. Do not add any explanations, quotation marks, or other text. Just the single word.
"""
    logging.info(
        f"Asking Gemini if book '{book_title}' is non-fiction for character list decision..."
    )

    # Create a cache prefix for this specific query
    safe_title_prefix = sanitize_filename(book_title, 30)
    cache_prefix = f"is_non_fiction_{safe_title_prefix}"

    response_text = call_llm_api(prompt, config, cache_prefix=cache_prefix)

    determined_is_non_fiction = False  # Default to fiction/indeterminate if API fails or gives unexpected response
    if response_text:
        answer = response_text.strip().lower()
        if answer == "yes":
            logging.info(f"Gemini indicates book '{book_title}' is likely non-fiction.")
            determined_is_non_fiction = True
        elif answer == "no":
            logging.info(
                f"Gemini indicates book '{book_title}' is likely fiction or its type is indeterminate."
            )
            determined_is_non_fiction = False
        else:
            logging.warning(
                f"Unexpected response from Gemini for non-fiction check ('{response_text}'). "
                "Treating as fiction/indeterminate."
            )
            # determined_is_non_fiction remains False
    else:
        logging.error(
            f"Failed to get response from Gemini for non-fiction check for '{book_title}'. "
            "Treating as fiction/indeterminate."
        )
        # determined_is_non_fiction remains False

    # Cache the 'is_fiction' status
    # If determined_is_non_fiction is True (book is non-fiction), then gen_params['is_fiction'] should be False.
    # If determined_is_non_fiction is False (book is fiction/indeterminate), then gen_params['is_fiction'] should be True.
    gen_params["is_fiction"] = not determined_is_non_fiction
    logging.info(
        f"Stored 'is_fiction': {gen_params['is_fiction']} in generation_params for '{book_title}'."
    )

    return determined_is_non_fiction


def generate_character_list(config, book_title):
    """
    Generates a list of character names and descriptions based on book details,
    if enabled in the config or if the book is determined to be fiction.
    """
    gen_params = config.get("generation_params", {})

    # Check for explicit override in config
    config_override = gen_params.get(
        "generate_character_list"
    )  # Can be True, False, or None

    should_generate = False  # Default assumption is no generation

    if config_override is False:
        # Config explicitly says NOT to generate
        logging.info("Character list generation explicitly disabled by config.")
        should_generate = False
    elif config_override is True:
        # Config explicitly says to generate (override fiction check)
        logging.info(
            "Character list generation explicitly enabled by config (overriding fiction check)."
        )
        should_generate = True
    else:  # config_override is None (key not present)
        # Fall back to the fiction/non-fiction check
        logging.info(
            "Character list generation not explicitly set in config. Using fiction/non-fiction check."
        )
        should_generate = not is_book_non_fiction(config, book_title)

    # If we decided not to generate, set the config key to None and return
    if not should_generate:
        gen_params["character_list"] = None  # Ensure key exists
        return None

    logging.info("Attempting to generate character list...")

    main_topic = gen_params.get("main_topic", "[No Main Topic Provided]")
    setting = gen_params.get("setting", "[No Setting Provided]")
    key_concepts = gen_params.get("key_concepts", [])

    # Critical prerequisites for character generation prompt
    if (
        main_topic == "[No Main Topic Provided]"
        or setting == "[No Setting Provided]"
        or not book_title  # book_title is a direct argument
    ):
        logging.error(
            "Cannot generate character list: 'main_topic', 'setting', or 'book_title' is missing/invalid. Skipping."
        )
        gen_params["character_list"] = None
        return None

    concepts_str = ", ".join(key_concepts) if key_concepts else "[No specific concepts]"

    prompt = f"""
Based on the book titled '{book_title}', which has the main topic '{main_topic}',
a setting described as: "{setting}", and key concepts including: {concepts_str}.

Generate a long list of characters that might appear in such a book.
For each character, provide their full name and a brief description
of their role, personality, or significance within the context of the topic and setting.

Format the output as a Markdown bulleted list. Each character should be an item.
Start the item with the character's name in bold, followed by a colon, and then the description.

Example:
*   **Character Name One:** A brief description of this character's role or significance.
*   **Another Character:** Their description and connection to the concepts.

Provide *only* the Markdown list of characters. Do not add introductory text like "Here is the character list:".
Output in British English.
"""

    character_list_text = call_llm_api(prompt, config, cache_prefix="character_list")

    if character_list_text:
        cleaned_text = character_list_text.strip()
        # Basic parsing: Split into lines and try to extract name/description
        characters = []
        for line in cleaned_text.split("\n"):
            line = line.strip()
            # Regex to capture bold name and the rest of the description
            match = re.match(r"^\*\s*\*\*(.*?)\*\*:\s*(.*)", line)
            if match:
                name = match.group(1).strip()
                description = match.group(2).strip()
                if name and description:
                    characters.append({"name": name, "description": description})
            elif line.startswith("* "):  # Handle cases where bolding might fail
                # Try a simpler split if regex fails but it looks like a list item
                parts = line[2:].split(":", 1)
                if len(parts) == 2 and parts[0].strip():
                    name = parts[0].strip()
                    description = parts[1].strip()
                    characters.append({"name": name, "description": description})

        if characters:
            logging.info(
                f"Successfully generated and parsed {len(characters)} characters."
            )
            # Store the parsed list in config
            gen_params["character_list"] = characters

            if characters:  # Check if the list is not empty
                character_log_details = "\n".join(
                    f"- {char.get('name', 'Unnamed')}: {char.get('description', 'No description')}"
                    for char in characters
                )
                logging.info(f"Generated Characters:\n{character_log_details}")

            return characters
        else:
            logging.warning(
                f"Could not parse character list from API response. Response:\n{cleaned_text}"
            )
            gen_params["character_list"] = None
            return None
    else:
        logging.error("Failed to generate character list via API.")
        gen_params["character_list"] = None
        return None


# --- Helper function to format character list for prompts ---
def format_character_list_for_prompt(character_list):
    """Formats the character list into a string suitable for API prompts."""
    if not character_list or not isinstance(character_list, list):
        return ""  # Return empty string if no characters

    formatted_items = []
    for char in character_list:
        if isinstance(char, dict) and "name" in char and "description" in char:
            formatted_items.append(f"- {char['name']}: {char['description']}")
        # Add handling for other potential formats if needed

    if not formatted_items:
        return ""

    return "Potential characters:\n" + "\n".join(formatted_items) + "\n"


def generate_chapter_outline(config, character_context=""):
    """Generates a list of chapter titles."""
    logging.info("Generating chapter outline...")
    length_modifier = (
        config.get("generation_params", {}).get("length_modifier", "").strip()
    )
    is_fiction = config.get("generation_params", {}).get("is_fiction", False)

    # Prepare the length_modifier part with a trailing space if it exists
    # e.g., "very " or ""
    actual_length_modifier_segment = f"{length_modifier} " if length_modifier else ""

    if not is_fiction:
        list_description_for_prompt = f"{actual_length_modifier_segment}short list"
    else:
        list_description_for_prompt = f"{actual_length_modifier_segment}long list"

    prompt = f"""Generate a {list_description_for_prompt} of chapter titles for a
book about '{config['generation_params']['main_topic']}'. The setting of the book
is described as: {config['generation_params']['setting']}.
Key concepts include: {', '.join(config['generation_params']['key_concepts'])}.
{character_context}
The chapters should logically progress through the topic, potentially involving the characters.
Ensure the list of chapter titles is appropriate for the type of book (fiction/non-fiction).
Format the output as a numbered list, with each title on a new line.
Start numbering from 1. Example:
1. Chapter Title One
2. Chapter Title Two
Do not use font formatting (e.g. bold, italics and etc) in the chapter title.
Do not generate two-part titles. The generated titles must not contain subtitles.
The chapter titles must not contain these punctuations: '-' or ':'.
Do not add introductory text.
Output in British English."""
    outline_text = call_llm_api(prompt, config, cache_prefix="chapter_outline")
    if outline_text:
        chapter_titles = []
        # Improved parsing to handle potential variations
        for line in outline_text.strip().split("\n"):
            line = line.strip()
            match = re.match(r"^\d+\.\s*(.*)", line)
            if match:
                title = match.group(1).strip()
                if title:  # Ensure title is not empty
                    chapter_titles.append(title)
        return chapter_titles
    else:
        logging.error("Failed to generate chapter outline via API. Using fallback.")
        return [
            f"Chapter {i+1}: Placeholder Title"
            for i in range(config["generation_params"]["num_chapters_fallback"])
        ]


def generate_chapter_summary(
    config,
    chapter_title,
    writing_tone,
    previous_summaries_context="",
    character_context="",
):
    """
    Generates a brief summary for a given chapter title, considering previous summaries.

    Args:
        config (dict): The application configuration.
        chapter_title (str): The title of the chapter to summarize.
        writing_tone (str): The desired writing tone.
        previous_summaries_context (str, optional): A string containing summaries
                                                    of preceding chapters. Defaults to "".
    """
    logging.info(f"Generating summary for chapter: '{chapter_title}'...")

    # --- Build the prompt ---
    prompt_parts = [
        f"Write a short and concise summary for the chapter titled '{chapter_title}'.",
        f"This chapter is part of a book about '{config['generation_params']['main_topic']}'.",
        f"The setting of the book is described as: {config['generation_params']['setting']}.",
        f"Key concepts: {', '.join(config['generation_params']['key_concepts'])}.",
        f"{character_context}",
    ]

    # Conditionally add context about previous chapters
    if previous_summaries_context:
        # Add the context clearly separated
        prompt_parts.append(
            f"\nContext: The summaries of the preceding chapters are:\n{previous_summaries_context}\n"
        )
        # Add instruction to avoid repetition
        prompt_parts.append(
            """Based on the preceding chapter summaries provided above, ensure 
this new summary is distinct and logically follows or contrasts with them, 
avoiding unnecessary repetition of themes or information already covered."""
        )

    # Prepare character-specific instruction
    character_consideration_text = ""
    if character_context:
        character_consideration_text = "Consider how the characters might be involved or relevant to this chapter's summary.\n"

    # Add remaining instructions
    prompt_parts.extend(
        [
            f"""\nMaintain a tone that is {writing_tone}.
{character_consideration_text}Output only the summary text for the current chapter ('{chapter_title}').
Do not add introductory text like 'This chapter summary is:'.
Output in British English."""
        ]
    )
    prompt = "\n".join(prompt_parts)
    logging.debug(
        f"Prompt for chapter summary '{chapter_title}':\n{prompt}"
    )  # Log the full prompt for debugging if needed
    # --- End prompt building ---

    summary_cache_prefix = f"summary_{sanitize_filename(chapter_title, max_length=40)}"
    summary = call_llm_api(prompt, config, cache_prefix=summary_cache_prefix)

    if summary:
        cleaned_summary = summary.strip()
        if cleaned_summary:
            logging.info(
                f"Successfully generated summary for '{chapter_title}':\n'{cleaned_summary}'"
            )
            return cleaned_summary
        else:
            logging.warning(
                f"Generated summary for '{chapter_title}' was empty after cleaning."
            )
            # Return placeholder if summary is empty after cleaning
            return f"Placeholder summary for chapter '{chapter_title}' focusing on {config['generation_params']['main_topic']}."
    else:
        logging.warning(
            f"Failed to generate summary for chapter '{chapter_title}'. Using placeholder."
        )
        return f"Placeholder summary for chapter '{chapter_title}' focusing on {config['generation_params']['main_topic']}."


def generate_section_titles(
    config,
    chapter_title,
    chapter_summary,
    all_chapter_titles,
    all_chapter_summaries,
    character_context="",
):
    """
    Generates a list of section titles for a given chapter, using its summary
    and considering the context of other chapters to avoid repetition.
    """
    logging.info(
        f"Generating section titles for chapter: '{chapter_title}' (using summary and full book context)..."
    )

    # --- Prepare Context for Prompt ---
    # Format all chapter titles
    all_titles_context = "\n".join(
        f"- {idx+1}. {title}" for idx, title in enumerate(all_chapter_titles)
    )

    # Format all chapter summaries (excluding the current one for brevity,
    # as it's provided separately)
    all_summaries_context_parts = []
    for idx, title in enumerate(all_chapter_titles):
        if title != chapter_title:  # Exclude current chapter's summary from this list
            summary = all_chapter_summaries.get(title, "[Summary not available]")
            all_summaries_context_parts.append(f"- Chapter '{title}': {summary}")
    all_summaries_context = "\n".join(all_summaries_context_parts)
    # --- End Context Preparation ---

    length_modifier = (
        config.get("generation_params", {}).get("length_modifier", "").strip()
    )

    prompt = f"""
Context for the entire book:
Main Topic: '{config['generation_params']['main_topic']}'
Setting: {config['generation_params']['setting']}
Key Concepts: {', '.join(config['generation_params']['key_concepts'])}
{character_context}

Full Chapter Outline:
{all_titles_context}

Summaries of OTHER chapters (for context on what's covered elsewhere):
{all_summaries_context}

---
Task:
Generate a {length_modifier} short list of relevant section titles specifically for the
chapter titled '{chapter_title}'.This chapter's specific summary is:
"{chapter_summary}"

Instructions:
- The section titles should logically break down the chapter's topic as 
described in *its specific summary*.
- Consider how the characters might relate to these sections.
- Ensure the generated section titles are distinct and avoid significant 
overlap with topics clearly covered in the *summaries of other chapters* 
provided above or topics strongly implied by the *titles of other chapters*.
- Format the output as a numbered list, with each title on a new line 
(e.g., 1. Section Title One).
- Do not generate two-part section titles. The generated section titles must
not contain subtitles.
- The section titles must not contain these punctuations: '-' or ':'.
- Do not use font formatting (e.g., bold, italics) in the section titles.
- Output *only* the numbered list of section titles. Do not add introductory text.
- Output in British English.
"""

    section_titles_cache_prefix = (
        f"section_titles_{sanitize_filename(chapter_title, max_length=40)}"
    )
    titles_text = call_llm_api(prompt, config, cache_prefix=section_titles_cache_prefix)
    num_chapter_fallback = config["generation_params"]["num_chapter_fallback"]
    if titles_text:
        section_titles = []
        # Improved parsing
        for line in titles_text.strip().split("\n"):
            line = line.strip()
            match = re.match(r"^\d+\.\s*(.*)", line)
            if match:
                title = match.group(1).strip()
                if title:
                    section_titles.append(title)

        # Use generated titles if we got *any*, otherwise fallback
        if section_titles:
            logging.info(
                f"Successfully generated {len(section_titles)} section titles for '{chapter_title}' considering context."
            )
            return section_titles
        else:
            logging.warning(
                f"Could not parse section titles for '{chapter_title}' from API response or response was empty. Using placeholders. Response:\n{titles_text}"
            )
            return [
                f"Section {i+1}: Placeholder Title" for i in range(num_chapter_fallback)
            ]
    else:
        logging.error(
            f"Failed to generate section titles for '{chapter_title}' via API. Using placeholders."
        )
        return [
            f"Section {i+1}: Placeholder Title" for i in range(num_chapter_fallback)
        ]


def generate_section_content(
    config,
    chapter_title,
    section_title,
    section_num,
    total_sections,
    chapter_summary,
    writing_tone,
    character_context="",
):
    """Generates content for a single section using Markdown, asking AI to use
    LaTeX math and avoid sub-headings."""
    logging.info(
        f"Generating content for: Chapter '{chapter_title}' -> Section {section_num}/{total_sections}: '{section_title}' (using summary)"
    )

    prompt = f"""
Context:
- Book Main Topic: '{config['generation_params']['main_topic']}'
- Setting: {config['generation_params']['setting']}
- Key Concepts: {', '.join(config['generation_params']['key_concepts'])}
{"-" if character_context else ""}{character_context}

---
Current Task Context:
- Current Chapter Title: '{chapter_title}'
- Current Chapter Summary: "{chapter_summary}"
- Current Section Title: '{section_title}'
- Current Section Number: {section_num} of {total_sections}
- Desired Writing Tone: {writing_tone}

Task:
Write a detailed section for the book described above, focusing specifically on 
the topic defined by the section title ('{section_title}'). Ensure the content 
fits logically within the context provided by the current chapter summary.

Instructions:
- Write approximately 2000 words for this section.
- Output *only* the text content for this section.
- Do *not* include the main chapter title or the section title in the output 
itself. Start directly with the section's content.
- Format the output using standard Markdown (paragraphs, lists, bold, italics,
tables).
- CRITICAL: Ensure all bulleted or numbered lists are preceded by a blank line in the 
Markdown output.
- Ensure paragraphs are separated by double line breaks in the Markdown source.
- Do *not* include any Markdown sub-headings (like ## Heading Level 2 or 
### Heading Level 3).
- If mathematical equations are necessary, format them using standard LaTeX 
syntax: use $...$ for inline math and $$...$$ for display math.
- Write the entire output in British English.
"""
    # --- Create the cache prefix ---
    # Sanitize chapter and section titles and combine them, limit length
    safe_chapter_title = sanitize_filename(chapter_title, max_length=30)
    safe_section_title = sanitize_filename(section_title, max_length=30)
    # Use a clear separator like '__' which is less common in titles
    cache_prefix_str = f"content_{safe_chapter_title}__{safe_section_title}"
    # --- End cache prefix creation ---

    # --- Call API with the prefix ---
    content = call_llm_api(prompt, config, cache_prefix=cache_prefix_str)
    # --- End API call ---

    return (
        content
        if content
        else f"**Content generation failed for Chapter '{chapter_title}', Section '{section_title}'.**"
    )


def generate_front_matter(
    config, book_title, author_name, writing_tone, summary_context
):
    """Generates front matter elements, including the subtitle."""
    logging.info("Generating front matter...")
    front_matter = {}

    # Subtitle generation already calls call_llm_api with its own prefix
    book_subtitle = generate_book_subtitle(config, book_title, summary_context)

    front_matter["title_page"] = {
        "title": book_title,
        "subtitle": book_subtitle,
        "author": author_name,
    }

    current_year = time.strftime("%Y")
    front_matter[
        "copyright_page"
    ] = f"""
Copyright © {current_year} by {author_name}



All rights reserved.



No part of this publication may be reproduced, distributed, or transmitted in
any form or by any means, including photocopying, recording, or other electronic
or mechanical methods, without the prior written permission of the publisher,
except in the case of brief quotations embodied in critical reviews and certain
other non-commercial uses permitted by copyright law. For permission requests,
contact {author_name}.



The story, all names, characters, and incidents portrayed in this production
are fictitious. No identification with actual persons (living or deceased),
places, buildings, and products are intended or should be inferred.



This publication is designed to provide accurate and authoritative information
in regard to the subject matter covered. It is sold with the understanding that
neither the author nor the publisher is engaged in rendering legal, investment,
accounting or other professional services. While the publisher and author have
used their best efforts in preparing this book, they make no representations or
warranties with respect to the accuracy or completeness of the contents of this
book and specifically disclaim any implied warranties of merchantability or
fitness for a particular purpose. No warranty may be created or extended by
sales representatives or written sales materials. The advice and strategies
contained herein may not be suitable for your situation. You should consult with
a professional when appropriate. Neither the publisher nor the author shall be
liable for any loss of profit or any other commercial damages, including but not
limited to special, incidental, consequential, personal, or other damages.
""".strip()

    common_prompt_base = f"""
for the book '{book_title}' about {config['generation_params']['main_topic']}, with the setting:
{config['generation_params']['setting']}. {summary_context}
Maintain a tone that is {writing_tone}. The author of this book is {author_name}.
Output *only* the text content for this section. Do not add introductory text.
Output in British English.
"""
    fm_elements_prompts = {
        "Dedication": f"Write an inspiring dedication {common_prompt_base}",
        "Foreword": f"Write a Foreword by a fictional expert relevant {common_prompt_base}. Make sure this fictional expert provides their name and credential at the end. Discuss the book's significance or context.",
        "Preface": f"Write a Preface {common_prompt_base}.  {author_name} explains their motivation or the book's scope.",
        "Acknowledgements": f"Write an Acknowledgements {common_prompt_base}. {author_name} thanks individuals and groups who contributed.",
    }
    for element, prompt in fm_elements_prompts.items():
        logging.info(f"Generating {element}...")
        # Use the element name (lowercase) as the prefix
        cache_prefix_str = element.lower()
        content = call_llm_api(prompt, config, cache_prefix=cache_prefix_str)

        processed_content = f"[{element} content generation failed.]"
        if content:
            content = content.strip()
            if content:
                expected_title = element
                lines = content.splitlines()
                processed_lines = list(lines)

                first_line_index = -1
                for i, line in enumerate(lines):
                    if line.strip():
                        first_line_index = i
                        break

                if first_line_index != -1:
                    first_line_content = lines[first_line_index].strip()
                    cleaned_first_line = re.sub(
                        r"^\s*#+\s*", "", first_line_content
                    ).strip()

                    if cleaned_first_line.lower() == expected_title.lower():
                        logging.debug(
                            f"Found and removing title '{lines[first_line_index]}' from {element} content."
                        )
                        del processed_lines[first_line_index]
                        while processed_lines and not processed_lines[0].strip():
                            logging.debug(
                                f"Removing blank line after title in {element}."
                            )
                            del processed_lines[0]
                        processed_content = "\n".join(processed_lines).strip()
                    else:
                        logging.debug(
                            f"First line of {element} ('{cleaned_first_line}') did not match expected title '{expected_title}'. Keeping original."
                        )
                        processed_content = content
                else:
                    processed_content = ""
            else:
                processed_content = ""
        else:
            logging.warning(f"Failed to generate content for {element}.")
            # processed_content already holds the error message

        front_matter[element.lower()] = processed_content

    return front_matter


def generate_back_matter(
    config, book_title, author_name, author_gender, writing_tone, summary_context
):
    """Generates back matter elements."""
    logging.info("Generating back matter...")
    back_matter = {}
    common_prompt_base = f"""
of the book '{book_title}' about {config['generation_params']['main_topic']}, with the setting:
{config['generation_params']['setting']} {summary_context}
Maintain a tone that is {writing_tone}.
Do not add introductory text. Output *only* the text content for this section.
Output in British English."""
    bm_elements_prompts = {
        "Appendix": f"Write an Appendix containing supplementary material relevant to the book's topic {common_prompt_base}",
        "Glossary": f"Create a Glossary defining key terms found in the book {common_prompt_base}",
        "Bibliography": f"Create a Bibliography listing fictional or real sources relevant to the book's content {common_prompt_base}",
        "About the Author": f"Write an 'About the Author' section for {author_gender} author {author_name} {common_prompt_base}",
    }
    for element, prompt in bm_elements_prompts.items():
        key = element.lower().replace(" ", "_")
        # Use the key as the prefix
        cache_prefix_str = key
        content = call_llm_api(prompt, config, cache_prefix=cache_prefix_str)
        back_matter[key] = (
            content.strip() if content else f"[{element} content generation failed.]"
        )

    return back_matter


def generate_book_blurb(config, book_title, summary_context, writing_tone):
    """Generates a marketing blurb for the book."""
    logging.info(f"Generating marketing blurb for book: '{book_title}'...")

    main_topic = config["generation_params"]["main_topic"]
    setting = config["generation_params"]["setting"]
    key_concepts = ", ".join(config["generation_params"]["key_concepts"])

    prompt = f"""Write a compelling marketing blurb for a book titled '{book_title}'.
The main topic of the book is '{main_topic}'.
The setting of the book is described as: {setting}.
Key concepts include: {key_concepts}.
{summary_context}
The blurb should entice readers while accurately reflecting the book's content.
Maintain a tone that is {writing_tone}, but adapted for marketing purposes (e.g., more engaging, intriguing).
Output only the blurb text. Do not add introductory text.
Output in British English."""

    blurb = call_llm_api(prompt, config, cache_prefix="book_blurb")

    if blurb:
        cleaned_blurb = blurb.strip()
        if cleaned_blurb:
            logging.info("Successfully generated book blurb.")
            return cleaned_blurb
        else:
            logging.warning(
                "Generated book blurb was empty after cleaning. Using placeholder."
            )
            return f"Placeholder blurb for book '{book_title}'."
    else:
        logging.warning("Failed to generate book blurb. Using placeholder.")
        return f"Placeholder blurb for book '{book_title}'."


def generate_overall_summary(config, book_title, summary_context):
    """Generates a single overall book summary using chapter summaries."""
    logging.info(f"Generating overall book summary for: '{book_title}'...")

    if not summary_context or summary_context == "[No chapter summaries available]":
        logging.warning(
            "Cannot generate overall summary: No chapter summaries available."
        )
        return f"Placeholder overall summary for the book '{book_title}'."

    prompt = f"""Based *only* on the following chapter summaries for the book 
titled '{book_title}', write a short overall summary or abstract of 
the entire book.


Chapter Summaries:
{summary_context}

Output *only* the overall summary text. Do not add introductory text.
Output in British English."""

    overall_summary = call_llm_api(prompt, config, cache_prefix="overall_book_summary")

    if overall_summary:
        cleaned_summary = overall_summary.strip()
        if cleaned_summary:
            logging.info("Successfully generated overall book summary.")
            return cleaned_summary
        else:
            logging.warning(
                "Generated overall summary was empty after cleaning. Using placeholder."
            )
            return f"Placeholder overall summary for the book '{book_title}'."
    else:
        logging.warning(
            "Failed to generate overall book summary via API. Using placeholder."
        )
        return f"Placeholder overall summary for the book '{book_title}'."


def save_summary_to_markdown(book_title, overall_summary, output_dir):
    """Saves the overall book summary to a Markdown file."""
    if not overall_summary:
        logging.warning("No overall summary provided to save.")
        return

    filename_stem = sanitize_filename(book_title)
    if not filename_stem or filename_stem.startswith("sanitized_empty"):
        logging.error(
            f"Could not create a valid filename from title '{book_title}'. Skipping summary Markdown save."
        )
        return

    output_filename = output_dir / f"{filename_stem}_Summary.md"
    markdown_content = f"# {book_title}\n\n{overall_summary}\n"

    try:
        with open(output_filename, "w", encoding="utf-8") as f:
            f.write(markdown_content)
        logging.info(f"Overall book summary saved to Markdown: '{output_filename}'")
    except Exception as e:
        logging.error(
            f"Error saving overall summary to Markdown file '{output_filename}': {e}"
        )


# --- DOCX Processing Functions --- (No changes needed in these for caching)


# Helper function to apply formatting to a run
def apply_formatting(run, bold=False, italic=False):
    """Applies formatting to a run."""
    run.bold = bold
    run.italic = italic


# Recursive function to process node content
def process_node_content(
    node,
    paragraph,
    container,
    doc,
    config,
    usable_width_inches,
    equation_image_dir,
    is_bold=False,
    is_italic=False,
):
    """
    Processes an lxml node's content (text and children) recursively,
    adding formatted runs to the paragraph or delegating block elements.
    Inherits formatting state.
    """
    if node.text:
        run = paragraph.add_run(node.text)
        apply_formatting(run, is_bold, is_italic)

    block_tags = {
        "p",
        "ul",
        "ol",
        "table",
        "blockquote",
        "pre",
        "hr",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "div",
    }

    for child in node:
        try:
            child_html_snippet = html.tostring(
                child, encoding="unicode", pretty_print=False
            )[:150]
        except Exception:
            child_html_snippet = f"Cannot serialize child <{child.tag}>"
        logging.debug(
            f"Processing child: tag=<{child.tag}>, class='{child.get('class', '')}', "
            f"has_text='{bool(child.text)}', has_tail='{bool(child.tail)}', "
            f"html='{child_html_snippet}...'"
        )

        node_class = child.get("class", "")
        child_tag = child.tag

        new_bold = is_bold or child_tag in ["strong", "b"]
        new_italic = is_italic or child_tag in ["em", "i"]

        is_block = child_tag in block_tags and not (
            child_tag == "div" and "arithmatex" in node_class
        )

        if is_block:
            logging.debug(
                f"-> Encountered block tag <{child_tag}> within inline processing. Calling add_paragraph_from_html_node."
            )
            if child.tail:
                logging.debug(
                    f"   Processing tail of block <{child_tag}> in original paragraph: '{child.tail[:50]}...'"
                )
                run = paragraph.add_run(child.tail)
                apply_formatting(
                    run, is_bold, is_italic
                )  # Apply formatting of the original paragraph context

        elif child_tag == "br":
            run = paragraph.add_run()
            run.add_break()
            if child.tail:
                run = paragraph.add_run(child.tail)
                apply_formatting(run, is_bold, is_italic)

        elif child.tag == "span" and "arithmatex" in node_class:  # Inline Math
            latex_code_with_delimiters = child.text_content().strip()
            logging.debug(
                f"-> Handling Inline Math Span. Raw Content: {latex_code_with_delimiters[:50]}..."
            )

            if latex_code_with_delimiters:
                image_path = render_latex_to_image(
                    latex_code_with_delimiters,
                    equation_image_dir,
                    is_display_style=False,
                )
                if image_path:
                    try:
                        pic_run = paragraph.add_run()
                        font_size_pt = 12
                        if (
                            paragraph.style
                            and paragraph.style.font
                            and paragraph.style.font.size
                        ):
                            font_size_pt = paragraph.style.font.size.pt
                        elif paragraph.runs:
                            for prev_run in reversed(paragraph.runs[:-1]):
                                if prev_run.font and prev_run.font.size:
                                    font_size_pt = prev_run.font.size.pt
                                    logging.debug(
                                        f"Detected font size {font_size_pt}pt from previous run."
                                    )
                                    break
                            else:
                                logging.debug(
                                    "Could not detect font size from previous runs, using default 12pt."
                                )
                        else:
                            logging.debug(
                                "Paragraph has no style/runs with size, using default 12pt."
                            )

                        inline_multiplier = config.get("style_params", {}).get(
                            "inline_math_height_multiplier",
                            DEFAULT_INLINE_MATH_HEIGHT_MULTIPLIER,
                        )
                        calculated_height = Pt(font_size_pt * inline_multiplier)
                        logging.debug(
                            f"Adding picture {image_path} with calculated height {calculated_height}"
                        )
                        pic_run.add_picture(image_path, height=calculated_height)

                        rpr = pic_run._r.get_or_add_rPr()
                        position_element = OxmlElement("w:position")
                        default_offset = -4
                        vertical_offset_half_points = config.get(
                            "style_params", {}
                        ).get("inline_math_vertical_offset_half_points", default_offset)
                        try:
                            vertical_offset_half_points = int(
                                vertical_offset_half_points
                            )
                        except (ValueError, TypeError):
                            logging.warning(
                                f"Invalid value '{vertical_offset_half_points}' for 'inline_math_vertical_offset_half_points' in config. Using default {default_offset}."
                            )
                            vertical_offset_half_points = default_offset
                        position_element.set(
                            qn("w:val"), str(vertical_offset_half_points)
                        )
                        rpr.append(position_element)
                        logging.debug(
                            f"Applied vertical offset ({vertical_offset_half_points} half-points) from config to inline math image run."
                        )

                    except Exception as img_err:
                        logging.error(
                            f"Error adding inline math picture {image_path}: {img_err}"
                        )
                        err_run = paragraph.add_run(
                            f"[Err: Inline Math '{latex_code_with_delimiters[:20]}...']"
                        )
                        apply_formatting(err_run, is_bold, is_italic)
                else:
                    err_run = paragraph.add_run(
                        f"[Render Err: {latex_code_with_delimiters[:20]}...]"
                    )
                    apply_formatting(err_run, is_bold, is_italic)
            else:
                logging.warning("-> Found inline math span but it was empty.")
                err_run = paragraph.add_run("[Err: Empty Math Span]")
                apply_formatting(err_run, is_bold, is_italic)

            if child.tail:
                run = paragraph.add_run(child.tail)
                apply_formatting(run, is_bold, is_italic)

        else:  # General Recursion for other INLINE tags
            logging.debug(f"-> Recursing inline into child <{child.tag}>...")
            process_node_content(
                child,
                paragraph,
                container,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
                new_bold,
                new_italic,
            )
            if child.tail:
                logging.debug(
                    f"   Processing tail of inline <{child.tag}>: '{child.tail[:50]}...'"
                )
                run = paragraph.add_run(child.tail)
                apply_formatting(run, is_bold, is_italic)


# --- In function: process_mixed_content ---
def process_mixed_content(
    parent_node,
    paragraph,
    container,
    doc,
    config,
    usable_width_inches,
    equation_image_dir,
):
    """
    Starts the recursive processing of mixed content within an HTML node (like p, li, td).
    Adds formatted runs directly to the provided paragraph object or delegates block elements.
    """
    process_node_content(
        parent_node,
        paragraph,
        container,
        doc,
        config,
        usable_width_inches,
        equation_image_dir,
        is_bold=False,
        is_italic=False,
    )


def delete_paragraph(paragraph):
    """Helper function to delete a paragraph."""
    p = paragraph._element
    if p is not None and p.getparent() is not None:
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None


def add_html_content_to_cell(
    html_node, cell, doc, config, usable_width_inches, equation_image_dir
):
    """Adds content from an HTML node (like TD or TH) to a DOCX cell."""
    # Clear existing content (Word adds an empty paragraph by default)
    for p in list(cell.paragraphs):  # Iterate over a copy
        delete_paragraph(p)

    # Add a new paragraph to start processing content
    p = cell.add_paragraph()
    process_mixed_content(
        html_node, p, cell, doc, config, usable_width_inches, equation_image_dir
    )

    # Ensure cell isn't totally empty (Word requires at least one paragraph)
    if not cell.paragraphs or (
        len(cell.paragraphs) == 1 and not cell.paragraphs[0].runs
    ):
        # If processing resulted in no paragraphs or an empty one, ensure one exists.
        if not cell.paragraphs:
            cell.add_paragraph()  # Add a paragraph if none exist
        # If the only paragraph is empty, it's fine, Word needs it.


def add_paragraph_from_html_node(
    node,
    container,
    doc,
    config,
    usable_width_inches,
    equation_image_dir,
    list_level=0,
):
    """
    Processes lxml HTML block nodes and adds them to a python-docx container (doc, cell).
    Handles p, h1-h6, ul, ol, li, blockquote, pre, hr, table, and display math divs/spans within p.
    Delegates inline formatting and nested block handling to process_mixed_content.
    Tracks list nesting depth. Treats all lists (ol and ul) as bulleted lists.
    """
    text = (node.text or "").strip()
    is_handled = False
    node_class = node.get("class", "")

    # --- (Keep add_display_math_image function as is) ---
    def add_display_math_image(image_path, latex_code_for_log):
        # ... (implementation unchanged) ...
        nonlocal container, doc, config, usable_width_inches
        try:
            is_main_body = hasattr(container, "add_table")  # Heuristic check
            p = container.add_paragraph()
            run = p.add_run()

            img_width_inches = None
            img_height_inches = None
            try:
                with Image.open(image_path) as img:
                    width_px, height_px = img.size
                    dpi = 300  # Match render_latex_to_image DPI
                    img_width_inches = width_px / dpi
                    img_height_inches = height_px / dpi
                    logging.debug(
                        f"Image {image_path}: {width_px}x{height_px}px @{dpi}dpi -> {img_width_inches:.2f}x{img_height_inches:.2f} inches"
                    )
            except Exception as pil_err:
                logging.error(
                    f"PIL Error reading image {image_path}: {pil_err}. Cannot determine size for scaling."
                )
                display_height_inches = config.get("style_params", {}).get(
                    "display_math_height_inches", DEFAULT_DISPLAY_MATH_HEIGHT_INCHES
                )
                run.add_picture(image_path, height=Inches(display_height_inches))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                return

            if (
                is_main_body
                and img_width_inches is not None
                and usable_width_inches is not None
                and img_width_inches > usable_width_inches
            ):
                scale_factor = usable_width_inches / img_width_inches
                final_width_inches = usable_width_inches
                final_height_inches = img_height_inches * scale_factor
                logging.info(
                    f'Scaling display math image {image_path} from {img_width_inches:.2f}" to fit usable width {usable_width_inches:.2f}" (scale: {scale_factor:.2f})'
                )
                run.add_picture(
                    image_path,
                    width=Inches(final_width_inches),
                    height=Inches(final_height_inches),
                )
            elif img_width_inches is not None:
                logging.debug(
                    f'Adding display math image {image_path} ({img_width_inches:.2f}") with original size (fits or not in main body).'
                )
                run.add_picture(
                    image_path,
                    width=Inches(img_width_inches),
                    height=Inches(img_height_inches),
                )
            else:  # Fallback if PIL failed earlier
                display_height_inches = config.get("style_params", {}).get(
                    "display_math_height_inches", DEFAULT_DISPLAY_MATH_HEIGHT_INCHES
                )
                run.add_picture(image_path, height=Inches(display_height_inches))

            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            logging.debug("Display math picture added and centered.")

        except Exception as img_err:
            logging.error(f"Error adding display math picture {image_path}: {img_err}")
            if "p" not in locals():
                p = container.add_paragraph()  # Ensure p exists
            p.add_run(
                f"[Error adding display math image: {latex_code_for_log[:30]}...]"
            )
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Handle Display Math (Arithmatex Div) ---
    if node.tag == "div" and "arithmatex" in node_class:
        # ... (implementation unchanged) ...
        logging.debug(
            f"Processing display math DIV node: <{node.tag} class='{node_class}'>"
        )
        latex_code_with_delimiters = node.text_content().strip()
        if latex_code_with_delimiters:
            image_path = render_latex_to_image(
                latex_code_with_delimiters, equation_image_dir, is_display_style=True
            )
            if image_path:
                logging.debug(f"Successfully rendered display math to {image_path}")
                add_display_math_image(image_path, latex_code_with_delimiters)
            else:
                logging.warning(
                    f"Rendering failed for display math: {latex_code_with_delimiters[:50]}..."
                )
                if hasattr(container, "add_paragraph"):
                    p = container.add_paragraph(
                        f"[Render Err: Display Math {latex_code_with_delimiters[:30]}...]"
                    )
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            logging.warning("Found display math div but it was empty.")
        is_handled = True

    # --- Handle Paragraphs (<p>) ---
    elif node.tag == "p":
        # ... (implementation unchanged, including display math span detection) ...
        logging.debug(f"Processing <p> node.")
        children = list(node)
        # Check if <p> *only* contains a display math span (common Arithmatex output)
        if (
            not text  # No text directly in <p>
            and len(children) == 1
            and children[0].tag == "span"
            and "arithmatex" in children[0].get("class", "")
            and not (children[0].tail or "").strip()  # No tail text after span
        ):
            logging.debug(
                "Detected <p> containing only an arithmatex span. Treating as display math."
            )
            span_node = children[0]
            latex_code_with_delimiters = span_node.text_content().strip()
            if latex_code_with_delimiters:
                image_path = render_latex_to_image(
                    latex_code_with_delimiters,
                    equation_image_dir,
                    is_display_style=True,
                )
                if image_path:
                    add_display_math_image(image_path, latex_code_with_delimiters)
                else:
                    logging.warning(
                        f"Rendering failed for display math span in p: {latex_code_with_delimiters[:50]}..."
                    )
                    if hasattr(container, "add_paragraph"):
                        p = container.add_paragraph(
                            f"[Render Err: Display Math {latex_code_with_delimiters[:30]}...]"
                        )
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                logging.warning("Found display math span within p but it was empty.")
            is_handled = True

        # --- Default paragraph handling ---
        if not is_handled:
            logging.debug("  Processing <p> using default process_mixed_content.")
            current_paragraph = container.add_paragraph()
            process_mixed_content(
                node,
                current_paragraph,
                container,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
            )
            # Remove paragraph if it ended up empty after processing
            if not current_paragraph.text and not current_paragraph.runs:
                logging.debug("Removing empty paragraph added for <p>.")
                delete_paragraph(current_paragraph)

            is_handled = True

    # --- Handle Tables ---
    elif node.tag == "table":
        # ... (implementation unchanged) ...
        logging.debug("Processing table...")
        html_rows = node.xpath(".//tr")  # Get all rows in the table
        if not html_rows:
            logging.warning("Table tag found but contains no rows (tr). Skipping.")
        else:
            # Determine number of columns from the first row
            first_row_cells = html_rows[0].xpath("./th|./td")
            num_cols = len(first_row_cells)
            if num_cols == 0:
                logging.warning(
                    "Table's first row contains no cells (th/td). Skipping table."
                )
            else:
                # Add table to the container (doc or cell)
                docx_table = container.add_table(rows=0, cols=num_cols)
                docx_table.style = "Table Grid"  # Apply a basic style

                # Process each row
                for html_row in html_rows:
                    docx_row = docx_table.add_row()
                    html_cells = html_row.xpath("./th|./td")
                    # Process each cell in the row
                    for i, cell_node in enumerate(html_cells):
                        if (
                            i < num_cols
                        ):  # Avoid index errors if rows have varying cell counts
                            docx_cell = docx_row.cells[i]
                            # Use the dedicated function to populate the cell
                            add_html_content_to_cell(
                                cell_node,
                                docx_cell,
                                doc,
                                config,
                                usable_width_inches,
                                equation_image_dir,
                            )
                        else:
                            logging.warning(
                                f"Row has more cells ({len(html_cells)}) than table columns ({num_cols}). Ignoring extra cells."
                            )
        is_handled = True

    # --- Handle Headings (h1-h6) ---
    elif node.tag in ["h1", "h2", "h3", "h4", "h5", "h6"]:
        # ... (implementation unchanged) ...
        level = int(node.tag[1])
        style_name = f"Heading {level}"
        # Use default paragraph style if heading style doesn't exist
        style = (
            doc.styles[style_name] if style_name in doc.styles else doc.styles["Normal"]
        )
        p = container.add_paragraph(style=style)
        process_mixed_content(
            node, p, container, doc, config, usable_width_inches, equation_image_dir
        )
        is_handled = True

    # --- Handle Lists (ul, ol, li) ---
    elif node.tag in ["ul", "ol"]:
        # Keep track of the level for children of THIS list
        child_list_level = list_level + 1
        logging.debug(
            f"Processing <{node.tag}> at list_level {list_level}. Children will be level {child_list_level}."
        )
        for child_node in node:
            # Only process li, ul, ol directly. Ignore whitespace text nodes etc.
            if child_node.tag == "li":
                # Pass the PARENT list's level down to the li
                add_paragraph_from_html_node(
                    child_node,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                    list_level=list_level,  # Pass current level to li
                )
            elif child_node.tag in ["ul", "ol"]:
                # Handle potentially invalid nested lists directly under lists
                logging.warning(
                    f"Found nested <{child_node.tag}> directly inside <{node.tag}>. Processing recursively."
                )
                add_paragraph_from_html_node(
                    child_node,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                    list_level=child_list_level,  # Increment level here
                )
            # Ignore other tags or text directly within ul/ol for now, or log warnings
            elif (child_node.text or "").strip():
                logging.warning(
                    f"Ignoring text '{child_node.text.strip()[:50]}...' found directly inside <{node.tag}>."
                )

        is_handled = True

    elif node.tag == "li":
        # Determine the style based on the level of the list this li belongs to.
        # Level 0 -> "List Bullet", Level 1 -> "List Bullet 2", etc.
        current_level_index = list_level + 1  # 1-based index for style name
        style_suffix = f" {current_level_index}" if current_level_index > 1 else ""
        # Assuming bullet lists for now, adapt if numbered lists needed different base
        style_name = f"List Bullet{style_suffix}"

        if style_name not in doc.styles:
            logging.warning(
                f"Style '{style_name}' not found. Falling back to 'List Bullet' or 'Normal'."
            )
            # Fallback logic: try base list style, then Normal
            style_name = "List Bullet" if "List Bullet" in doc.styles else "Normal"

        style = doc.styles[style_name]
        logging.debug(
            f"Processing <li> at list_level {list_level} using style '{style.name}'."
        )

        # Create the paragraph for this list item's text/inline content
        # Check if the first child is a <p> tag (common markdown output)
        first_child_is_p = False
        direct_children = [child for child in node if isinstance(child.tag, str)]
        if (
            direct_children
            and direct_children[0].tag == "p"
            and not (node.text or "").strip()
        ):
            # If li starts directly with a <p>, use its content but apply list style
            p_node = direct_children[0]
            p = container.add_paragraph(style=style)
            process_mixed_content(
                p_node,
                p,
                container,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
            )
            # Remove the processed <p> node from children list to avoid double processing
            children_to_process = direct_children[1:]
            logging.debug("Processed <li> content starting with <p>.")
        else:
            # Process text/inline elements directly under <li>
            p = container.add_paragraph(style=style)
            process_mixed_content(
                node, p, container, doc, config, usable_width_inches, equation_image_dir
            )
            # We processed the whole node inline content, but need to handle block children (nested lists) separately
            children_to_process = (
                direct_children  # Re-evaluate children for nested lists
            )
            logging.debug("Processed <li> content directly.")

        # Now, specifically look for and handle nested lists *within* this <li>
        nested_list_found = False
        for (
            child_node
        ) in children_to_process:  # Use the adjusted list if first child was <p>
            if child_node.tag in ["ul", "ol"]:
                nested_list_found = True
                # Recursively call for the nested list, INCREMENTING the level
                nested_list_level = list_level + 1
                logging.debug(
                    f"  Found nested <{child_node.tag}> inside <li>. Processing at level {nested_list_level}."
                )
                add_paragraph_from_html_node(
                    child_node,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                    list_level=nested_list_level,  # Pass incremented level
                )
            # Note: We assume process_mixed_content handled inline tags and text already.
            # If there were other block elements inside <li> (besides <p> handled above),
            # they might need specific handling here too, but nested lists are primary.

        # Remove the list item paragraph ONLY if it's empty AND no nested list followed.
        # Check runs as well as text, because an image (like math) adds a run but no text.
        if not p.text.strip() and not p.runs and not nested_list_found:
            logging.debug(
                f"Removing empty paragraph potentially created for <li> at level {list_level}."
            )
            delete_paragraph(p)

        is_handled = True

    # --- Handle Blockquotes ---
    elif node.tag == "blockquote":
        logging.debug("Processing <blockquote> node.")
        style = doc.styles["Quote"] if "Quote" in doc.styles else doc.styles["Normal"]
        for child_node in node:
            if child_node.tag == "p":
                p = container.add_paragraph(style=style)
                process_mixed_content(
                    child_node,
                    p,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                )
                if not p.text and not p.runs:
                    logging.debug(
                        "Removing empty paragraph added for <p> inside <blockquote>."
                    )
                    delete_paragraph(p)
            elif child_node.tag is not None:
                logging.debug(
                    f"Handling non-<p> tag '{child_node.tag}' inside blockquote."
                )
                add_paragraph_from_html_node(
                    child_node,
                    container,
                    doc,
                    config,
                    usable_width_inches,
                    equation_image_dir,
                )
            elif (child_node.text or "").strip():
                logging.debug("Handling text node directly inside <blockquote>.")
                p = container.add_paragraph(child_node.text.strip(), style=style)
        is_handled = True

    # --- Handle Highlighted Code Blocks (div.highlight > pre > code) ---
    elif node.tag == "div" and "highlight" in node.get("class", "").split():
        # ... (implementation unchanged) ...
        logging.debug("Processing <div class='highlight'> node.")
        pre_node = node.find("pre")
        if pre_node is not None:
            code_node = pre_node.find("code")
            # Extract text from <code> if present, otherwise from <pre>
            full_text = (
                "".join(pre_node.itertext())
                if code_node is None
                else "".join(code_node.itertext())
            )
            if full_text:
                code_style_name = "CodeBlock"
                style_to_use = (
                    doc.styles[code_style_name]
                    if code_style_name in doc.styles
                    else doc.styles["Normal"]
                )
                logging.debug(
                    f"Applying style '{style_to_use.name}' to <div class='highlight'> content."
                )

                p = container.add_paragraph(full_text.strip("\n"), style=style_to_use)

                # Explicitly set font on runs as a fallback
                for run in p.runs:
                    if not run.font.name or run.font.name != "Courier New":
                        run.font.name = "Courier New"
                    if not run.font.size or run.font.size != Pt(10):
                        run.font.size = Pt(10)
                logging.debug(
                    "Ensured Courier New/10pt font on runs within the highlight code block."
                )
            else:
                logging.debug("Highlight div found, but contained no text.")
        else:
            logging.warning(
                "Found <div class='highlight'> but no <pre> tag inside. Skipping."
            )
        is_handled = True

    # --- Handle Preformatted Text (<pre>) ---
    elif node.tag == "pre":
        # ... (implementation unchanged) ...
        parent = node.getparent()
        if (
            parent is not None
            and parent.tag == "div"
            and "highlight" in parent.get("class", "").split()
        ):
            logging.debug(
                "Skipping <pre> inside already handled <div class='highlight'>."
            )
            is_handled = True
        else:
            logging.debug("Processing plain <pre> node (not inside highlight div).")
            code_node = node.find("code")
            full_text = (
                "".join(node.itertext())
                if code_node is None
                else "".join(code_node.itertext())
            )
            if full_text:
                code_style_name = "CodeBlock"
                style_to_use = (
                    doc.styles[code_style_name]
                    if code_style_name in doc.styles
                    else doc.styles["Normal"]
                )
                logging.debug(f"Applying style '{style_to_use.name}' to <pre> content.")

                p = container.add_paragraph(full_text.strip("\n"), style=style_to_use)

                for run in p.runs:
                    if not run.font.name or run.font.name != "Courier New":
                        run.font.name = "Courier New"
                    if not run.font.size or run.font.size != Pt(10):
                        run.font.size = Pt(10)
                logging.debug(
                    "Ensured Courier New/10pt font on runs within the code block."
                )
            is_handled = True

    # --- Handle Horizontal Rule (<hr>) ---
    elif node.tag == "hr":
        # ... (implementation unchanged) ...
        logging.debug("Processing <hr> node.")
        p = container.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        pPr.insert_element_before(
            pBdr,
            "w:shd",
            "w:tabs",
            "w:suppressAutoHyphens",
            "w:kinsoku",
            "w:wordWrap",
            "w:overflowPunct",
            "w:topLinePunct",
            "w:autoSpaceDE",
            "w:autoSpaceDN",
            "w:bidi",
            "w:adjustRightInd",
            "w:snapToGrid",
            "w:spacing",
            "w:ind",
            "w:contextualSpacing",
            "w:mirrorIndents",
            "w:suppressOverlap",
            "w:jc",
            "w:textDirection",
            "w:textAlignment",
            "w:textboxTightWrap",
            "w:outlineLvl",
            "w:divId",
            "w:cnfStyle",
            "w:rPr",
            "w:sectPr",
            "w:pPrChange",
        )
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 3/4 pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        is_handled = True

    # --- Fallback for Unhandled Block Tags ---
    known_inline_or_handled = {
        "strong",
        "b",
        "em",
        "i",
        "span",
        "br",
        "a",
        "code",  # Common inline
        "p",
        "ul",
        "ol",
        "li",
        "table",
        "tr",
        "td",
        "th",  # Handled block
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",  # Handled block
        "blockquote",
        "pre",
        "hr",
        "div",  # Handled block (incl. arithmatex)
    }
    if (
        not is_handled
        and node.tag is not None  # Ensure it's a tag, not comment/text
        and node.tag not in known_inline_or_handled
    ):
        logging.warning(
            f"Unhandled block tag <{node.tag}> encountered. Attempting to process its content as plain text."
        )
        plain_text = "".join(node.itertext()).strip()
        if plain_text:
            container.add_paragraph(plain_text)
            logging.debug(
                f"Added text content of unhandled <{node.tag}>: '{plain_text[:100]}...'"
            )
        else:
            logging.debug(f"Unhandled block tag <{node.tag}> had no text content.")
        is_handled = True


def markdown_to_docx(
    markdown_text,
    container_obj,
    doc,
    config,
    usable_width_inches,
    equation_image_dir,
    context_label=None,
):
    """Converts Markdown to DOCX elements, using Arithmatex for LaTeX,
    and cleans up excessive line breaks. Handles nested lists."""
    # ... (keep markdown cleaning and HTML conversion as is) ...
    if not markdown_text:
        logging.debug("Markdown text is empty, skipping conversion.")
        return

    logging.debug("Starting Markdown to DOCX conversion (using Arithmatex)...")

    # --- Pre-processing: Clean up excessive newlines ---
    cleaned_markdown = re.sub(r"\n{3,}", "\n\n", markdown_text.strip())
    if cleaned_markdown != markdown_text.strip():
        logging.debug("Cleaned excessive newlines from Markdown content.")
    # --- End Pre-processing ---

    try:
        # Configure Markdown extensions
        extensions = [
            "extra",  # Includes tables, footnotes, abbreviations, etc.
            "sane_lists",
            "fenced_code",
            "pymdownx.arithmatex",  # For LaTeX math $$...$$ and $...$
            "pymdownx.superfences",  # Improved fenced code blocks
            "pymdownx.details",
            "pymdownx.mark",
        ]

        extension_configs = {"pymdownx.arithmatex": {"generic": True}}
        md_converter = markdown.Markdown(
            extensions=extensions, extension_configs=extension_configs
        )
        html_content = md_converter.convert(cleaned_markdown)
        logging.debug(f"Generated HTML (first 500 chars): {html_content[:500]}...")

        # --- Save intermediate HTML for debugging ---
        if config.get("debug_options", {}).get("save_intermediate_html", False):
            # ... (HTML saving logic unchanged) ...
            debug_html_dir = (
                pathlib.Path(config.get("cache_dir", "api_cache")) / "debug_html"
            )
            debug_html_dir.mkdir(parents=True, exist_ok=True)
            timestamp = time.strftime("%Y%m%d_%H%M%S")
            content_hash = hashlib.sha1(cleaned_markdown.encode("utf-8")).hexdigest()[
                :8
            ]
            if context_label:
                context_hint = sanitize_filename(context_label, max_length=80)
            else:
                context_hint = "unknown_context"

            html_filename = (
                debug_html_dir
                / f"md_to_html_{context_hint}_{timestamp}_{content_hash}.html"
            )
            try:
                with open(html_filename, "w", encoding="utf-8") as f_html:
                    f_html.write(
                        "<!DOCTYPE html>\n<html>\n<head><meta charset='UTF-8'>"
                    )
                    f_html.write(
                        "<style>.arithmatex { border: 1px dotted blue; padding: 2px; }</style>"
                    )
                    f_html.write("</head>\n<body>\n")
                    f_html.write(html_content)
                    f_html.write("\n</body>\n</html>")
                logging.info(f"Saved intermediate HTML to: {html_filename}")
            except Exception as e_save:
                logging.error(
                    f"Could not save intermediate HTML to {html_filename}: {e_save}"
                )
        # --- End intermediate HTML saving ---

        # --- HTML Cleaning Step ---
        parser = html.HTMLParser(encoding="utf-8")
        # Wrap in a div to ensure a single root for parsing potentially fragmented HTML
        html_wrapper_str = f"<div>{html_content}</div>"
        try:
            # Use utf-8 encoding for parsing
            html_tree_root = html.fromstring(
                html_wrapper_str.encode("utf-8"), parser=parser
            )
        except UnicodeDecodeError:
            # Fallback if utf-8 fails (less common but possible)
            logging.warning(
                "UTF-8 decoding failed for HTML string, trying 'latin-1' fallback."
            )
            html_tree_root = html.fromstring(
                html_wrapper_str.encode("latin-1"), parser=parser
            )

        # Find <p> tags that are either completely empty or contain only <br> tags and whitespace
        paragraphs_to_remove = []
        # Iterate through all <p> tags in the parsed HTML tree
        for p_tag in html_tree_root.xpath(".//p"):
            # Get the text content of the <p> tag, stripping leading/trailing whitespace
            text_content = p_tag.text_content().strip()
            # Get all direct children elements of the <p> tag
            children = p_tag.getchildren()

            # Check condition 1: Is the paragraph completely empty (no text, no children)?
            is_completely_empty = not text_content and not children

            # Check condition 2: Does the paragraph contain ONLY <br> tags (and whitespace)?
            only_br_children = False
            if not text_content and children:  # Only check children if there's no text
                all_children_are_br = True
                for child in children:
                    # If any child is not a <br> tag, this condition is false
                    if child.tag != "br":
                        all_children_are_br = False
                        break
                only_br_children = all_children_are_br  # True if all children were <br>

            # If either condition is met, schedule the paragraph for removal
            if is_completely_empty or only_br_children:
                logging.debug(
                    f"Found <p> tag {'empty' if is_completely_empty else 'containing only <br> tags'}. Scheduling for removal. "
                    f"HTML snippet: {html.tostring(p_tag, encoding='unicode', pretty_print=False)[:100]}"
                )
                paragraphs_to_remove.append(p_tag)

        # Remove the identified paragraphs
        if paragraphs_to_remove:
            logging.info(
                f"Removing {len(paragraphs_to_remove)} empty or <br>-only <p> tags."
            )
            for p_tag in paragraphs_to_remove:
                parent = p_tag.getparent()
                if parent is not None:
                    # Preserve tail text if it exists, attaching it to the previous sibling or parent text
                    if p_tag.tail and p_tag.tail.strip():
                        previous_sibling = p_tag.getprevious()
                        if previous_sibling is not None:
                            # Append tail to the previous sibling's tail
                            previous_sibling.tail = (
                                previous_sibling.tail or ""
                            ) + p_tag.tail
                        else:
                            # Append tail to the parent's text if no previous sibling
                            parent.text = (parent.text or "") + p_tag.tail
                    # Remove the paragraph tag itself
                    parent.remove(p_tag)
        # --- End HTML Cleaning Step ---

        # Process each *remaining* element under the root div
        for element in html_tree_root:
            add_paragraph_from_html_node(
                element,
                container_obj,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
            )

    except ImportError as ie:
        # ... (error handling unchanged) ...
        logging.error(
            f"Markdown extension import error: {ie}. Ensure required libraries (e.g., pymdown-extensions) are installed.",
            exc_info=True,
        )
        container_obj.add_paragraph(f"[Error: Missing Markdown extension - {ie}]")
        container_obj.add_paragraph(cleaned_markdown)
    except Exception as e:
        # ... (error handling unchanged) ...
        logging.error(f"Error converting Markdown/HTML to DOCX: {e}", exc_info=True)
        container_obj.add_paragraph(f"[Error processing content: {e}]")
        container_obj.add_paragraph("--- Raw Markdown Fallback ---")
        container_obj.add_paragraph(cleaned_markdown)
        container_obj.add_paragraph("--- End Raw Markdown ---")

    logging.debug("Finished Markdown to DOCX conversion.")


def set_page_numbering(section, format_code, start_number=None):
    """Adds page numbering to the footer of a given section."""
    # Ensure footer exists (it might not by default)
    if section.footer is None:
        section.footer  # Accessing it creates it
        logging.debug("Created footer for section.")

    footer = section.footer

    # Unlink footer if setting a start number and it was linked
    if footer.is_linked_to_previous and start_number is not None:
        footer.is_linked_to_previous = False
        logging.debug(f"Unlinking footer for section starting page {start_number}")

    # Clear existing footer content (if any) before adding page number field
    if footer.paragraphs:
        pf = footer.paragraphs[0]
        pf.clear()  # Clear runs from the first paragraph
        # Remove other paragraphs if they exist
        for p in footer.paragraphs[1:]:
            delete_paragraph(p)
    else:
        # Add a paragraph if none exist
        pf = footer.add_paragraph()

    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run_begin = pf.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_char_begin)

    run_instr = pf.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    run_instr._r.append(instr_text)

    run_sep = pf.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_char_separate)

    # Optional: Add a run here for the actual number display if needed,
    # but Word usually handles this automatically with the fields.

    run_end = pf.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_char_end)
    # --- End PAGE field ---

    # --- Set page number type in section properties ---
    sectPr = section._sectPr
    # Remove existing pgNumType if it exists to avoid conflicts
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is not None:
        sectPr.remove(pgNumType)

    # Create and append new pgNumType element
    pgNumType = OxmlElement("w:pgNumType")
    sectPr.append(pgNumType)  # Append to the end or specific location if needed

    # Set format
    pgNumType.set(qn("w:fmt"), format_code)

    # Set start number if provided
    if start_number is not None:
        pgNumType.set(qn("w:start"), str(start_number))
        logging.debug(
            f"Set page numbering: format='{format_code}', start={start_number}"
        )
    else:
        # Ensure start attribute is removed if not specified (to allow continuation)
        if qn("w:start") in pgNumType.attrib:
            del pgNumType.attrib[qn("w:start")]
        logging.debug(
            f"Set page numbering: format='{format_code}', continuing from previous section."
        )


def assemble_docx(
    config,
    front_matter,
    body_matter,
    back_matter,
    book_title,
    equation_image_dir,
    output_dir,
):
    """Assembles the main book DOCX file with complex page numbering and MathML/OXML."""
    logging.info("Assembling main DOCX file...")
    filename_stem = sanitize_filename(book_title)
    # Construct the full output path using the provided directory
    output_filename = output_dir / f"{filename_stem}.docx"
    logging.info(f"Main book filename set to: '{output_filename}'")

    style_config = config.get("style_params", {})
    font_name = style_config.get("font_name", "Times New Roman")
    font_size = style_config.get("font_size", 12)
    page_preset = style_config.get("page_size_preset", "6x9")

    margin_config_mm = style_config.get("margins_mm", {})
    default_top_mm = 19  # Approx 0.75 inch
    default_bottom_mm = 19
    default_inside_mm = 19  # For gutter
    default_outside_mm = 13  # Approx 0.5 inch
    default_gutter_mm = 0  # Set gutter explicitly

    top_margin_mm = margin_config_mm.get("top", default_top_mm)
    bottom_margin_mm = margin_config_mm.get("bottom", default_bottom_mm)
    # 'left' in config maps to 'outside', 'right' maps to 'inside' for mirrored margins
    outside_margin_mm = margin_config_mm.get("left", default_outside_mm)
    inside_margin_mm = margin_config_mm.get("right", default_inside_mm)
    gutter_margin_mm = margin_config_mm.get("gutter", default_gutter_mm)

    doc = Document()

    # --- Basic Style Setup ---
    try:
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Ensure heading styles use the base font and add spacing
        for i in range(1, 7):
            heading_style_name = f"Heading {i}"
            if heading_style_name in doc.styles:
                h_style = doc.styles[heading_style_name]
                h_style.font.name = font_name
                # Add some default spacing (can be overridden in config later)
                if i == 1:
                    h_style.paragraph_format.space_before = Pt(18)
                    h_style.paragraph_format.space_after = Pt(6)
                elif i == 2:
                    h_style.paragraph_format.space_before = Pt(12)
                    h_style.paragraph_format.space_after = Pt(4)
                else:
                    h_style.paragraph_format.space_before = Pt(6)
                    h_style.paragraph_format.space_after = Pt(2)

        # Ensure Title/Subtitle styles use the base font (or define them)
        if "Title" not in doc.styles:
            title_style = doc.styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
            title_style.base_style = doc.styles["Normal"]
            title_style.font.name = font_name
            title_style.font.size = Pt(28)  # Example size
            title_style.font.bold = True
            title_style.paragraph_format.space_after = Pt(6)
            title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.styles["Title"].font.name = font_name
            doc.styles[
                "Title"
            ].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if "Subtitle" not in doc.styles:
            subtitle_style = doc.styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.base_style = doc.styles["Normal"]
            subtitle_style.font.name = font_name
            subtitle_style.font.size = Pt(16)  # Example size
            subtitle_style.font.italic = True
            subtitle_style.paragraph_format.space_after = Pt(18)
            subtitle_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            doc.styles["Subtitle"].font.name = font_name
            doc.styles[
                "Subtitle"
            ].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Ensure List Bullet exists (as before)
        if "List Bullet" not in doc.styles:
            lb_style = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
            lb_style.base_style = doc.styles["Normal"]
            # TODO: Define actual bullet point and indentation via Oxml or basic properties
            lb_style.paragraph_format.left_indent = Inches(0.25)  # Example
            lb_style.paragraph_format.first_line_indent = Inches(
                -0.25
            )  # Example hanging indent
        else:  # Ensure base style has some indent
            lb_style = doc.styles["List Bullet"]
            if lb_style.paragraph_format.left_indent is None:
                lb_style.paragraph_format.left_indent = Inches(0.25)
            if lb_style.paragraph_format.first_line_indent is None:
                lb_style.paragraph_format.first_line_indent = Inches(-0.25)

        # Define nested styles (add more as needed)
        for i in range(2, 5):  # Define List Bullet 2, 3, 4
            style_name = f"List Bullet {i}"
            base_style_name = (
                f"List Bullet {i-1}" if i > 2 else "List Bullet"
            )  # Base on previous level
            if style_name not in doc.styles:
                lb_nested_style = doc.styles.add_style(
                    style_name, WD_STYLE_TYPE.PARAGRAPH
                )
                # Base on previous level if possible, otherwise Normal
                base_style = (
                    doc.styles[base_style_name]
                    if base_style_name in doc.styles
                    else doc.styles["Normal"]
                )
                lb_nested_style.base_style = base_style
                # Increase indentation relative to base style or set absolute
                # Example: Add 0.25 inches per level
                indent_inches = (i - 1) * 0.35  # Adjust multiplier as needed
                lb_nested_style.paragraph_format.left_indent = Inches(indent_inches)
                # Keep hanging indent consistent or adjust if needed
                lb_nested_style.paragraph_format.first_line_indent = Inches(-0.25)
                logging.info(
                    f"Defined style '{style_name}' with left indent {indent_inches} inches."
                )
            else:
                # Optionally ensure indentation is correct on existing styles
                lb_nested_style = doc.styles[style_name]
                indent_inches = (i - 1) * 0.35
                if lb_nested_style.paragraph_format.left_indent != Inches(
                    indent_inches
                ):
                    logging.debug(f"Adjusting indent for existing style '{style_name}'")
                    lb_nested_style.paragraph_format.left_indent = Inches(indent_inches)
                    lb_nested_style.paragraph_format.first_line_indent = Inches(-0.25)

        # --- Define Code Block Style ---
        code_style_name = "CodeBlock"
        if code_style_name not in doc.styles:
            code_style = doc.styles.add_style(code_style_name, WD_STYLE_TYPE.PARAGRAPH)
            # Base on 'No Spacing' if it exists for minimal vertical space, else 'Normal'
            base_style_name = "No Spacing" if "No Spacing" in doc.styles else "Normal"
            code_style.base_style = doc.styles[base_style_name]
            code_style.font.name = "Courier New"  # Monospace font
            code_style.font.size = Pt(10)  # Slightly smaller size often looks good
            # Optional: Add indentation or borders
            # code_style.paragraph_format.left_indent = Inches(0.25)
            # Optional: Adjust spacing if needed (base style might handle it)
            # code_style.paragraph_format.space_before = Pt(6)
            # code_style.paragraph_format.space_after = Pt(6)
            logging.info(
                f"Defined '{code_style_name}' style based on '{base_style_name}'."
            )
        else:
            # Ensure existing style uses monospace font
            existing_code_style = doc.styles[code_style_name]
            existing_code_style.font.name = "Courier New"
            existing_code_style.font.size = Pt(10)
            logging.info(f"Ensured '{code_style_name}' style uses Courier New, 10pt.")
        # --- End Code Block Style ---

        logging.info("Styles configured.")
    except Exception as e:
        logging.error(f"Error setting up styles: {e}")

    # --- Page Setup (Initial Section - Section 0) ---
    section0 = doc.sections[0]
    page_width_mm = None

    if page_preset == "6x9":
        section0.page_width = Inches(6)
        section0.page_height = Inches(9)
        logging.info("Set page size to 6x9 inches.")
    elif page_preset == "A4":
        section0.page_width = Mm(210)
        section0.page_height = Mm(297)
        logging.info("Set page size to A4.")
    else:
        logging.warning(
            f"Unsupported page_size_preset '{page_preset}'. Using default Word size."
        )
        # Use default size implicitly

    if section0.page_width is not None:
        page_width_mm = section0.page_width / Mm(1)
        logging.info(f"Actual page width from section object: {page_width_mm:.2f} mm")
    else:
        logging.warning("Could not determine page width from section object.")

    # --- Apply Margins and Gutter (Mirrored) ---
    try:
        section0.top_margin = Mm(top_margin_mm)
        section0.bottom_margin = Mm(bottom_margin_mm)
        section0.left_margin = Mm(outside_margin_mm)  # Outside margin
        section0.right_margin = Mm(inside_margin_mm)  # Inside margin
        section0.gutter = Mm(gutter_margin_mm)
        # Enable mirrored margins for gutter to work correctly
        sectPr = section0._sectPr
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            pgMar.set(qn("w:mirrorMargins"), "true")  # Use 'true' or '1'
        logging.info(
            f"Set mirrored margins (mm): Top={top_margin_mm}, Bottom={bottom_margin_mm}, "
            f"Outside={outside_margin_mm}, Inside={inside_margin_mm}, Gutter={gutter_margin_mm}"
        )
    except ValueError as ve:
        logging.error(f"Invalid margin value provided: {ve}. Using Word defaults.")
    except Exception as e:
        logging.error(f"Error setting margins: {e}. Using Word defaults.")

    # --- Calculate Usable Width ---
    usable_width_mm = None
    usable_width_inches = None
    if page_width_mm is not None:
        # Usable width = Page Width - Outside Margin - Inside Margin - Gutter
        usable_width_mm = (
            page_width_mm - outside_margin_mm - inside_margin_mm - gutter_margin_mm
        )
        usable_width_inches = usable_width_mm / 25.4
        logging.info(
            f"Calculated usable page width: {usable_width_mm:.2f} mm ({usable_width_inches:.2f} inches)"
        )
    else:
        logging.error("Cannot calculate usable width because page width is unknown.")
    # --- End Calculate Usable Width ---

    # --- Section 0: Title Page ---
    logging.info("Adding Title Page (Section 0)...")
    if "title_page" in front_matter:
        tp_info = front_matter["title_page"]
        # Add space before title (adjust as needed)
        doc.add_paragraph().paragraph_format.space_before = Pt(72)

        title_p = doc.add_paragraph(tp_info["title"], style="Title")
        # Alignment is set in style definition now

        if tp_info.get("subtitle"):
            subtitle_p = doc.add_paragraph(tp_info["subtitle"], style="Subtitle")
            # Alignment and spacing set in style definition
        else:
            # Add extra space after title if no subtitle
            title_p.paragraph_format.space_after = Pt(
                24
            )  # Override style default if needed

        # Add Author with space before it
        author_p = doc.add_paragraph()  # Empty paragraph for spacing
        author_p.paragraph_format.space_before = Pt(36)  # Space before author
        author_p.add_run(f"By {tp_info['author']}")
        author_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Optional: Create and apply an 'Author' style

    # --- Section Break for Copyright Page (Starts Section 1) ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    section1 = doc.sections[1]
    # Copy page setup from section 0 to section 1
    section1.page_height = section0.page_height
    section1.page_width = section0.page_width
    section1.left_margin = section0.left_margin  # Outside
    section1.right_margin = section0.right_margin  # Inside
    section1.top_margin = section0.top_margin
    section1.bottom_margin = section0.bottom_margin
    section1.gutter = section0.gutter
    # Ensure mirrored margins are also copied
    sectPr1 = section1._sectPr
    pgMar1 = sectPr1.find(qn("w:pgMar"))
    if pgMar1 is not None:
        pgMar1.set(qn("w:mirrorMargins"), "true")
    # No page numbering for title (section 0) or copyright (section 1)

    # --- Section 1: Copyright Page ---
    logging.info("Adding Copyright Page (Section 1)...")
    if "copyright_page" in front_matter:
        cp_text = front_matter["copyright_page"]
        # Split into paragraphs based on double line breaks in the original string
        cp_paragraphs = re.split(r"\n\s*\n", cp_text)

        # Add space before the first paragraph
        first_cp_p = doc.add_paragraph()
        first_cp_p.paragraph_format.space_before = Pt(60)

        # FIX: Replace internal newlines with spaces before adding
        first_para_text = re.sub(r"\s*\n\s*", " ", cp_paragraphs[0].strip())
        first_cp_p.add_run(first_para_text)

        first_cp_p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Apply smaller font size
        for run in first_cp_p.runs:
            run.font.size = Pt(font_size - 2)  # Assuming font_size is defined

        # Add subsequent paragraphs
        for cp_para in cp_paragraphs[1:]:
            # FIX: Replace internal newlines with spaces before adding
            para_text = re.sub(r"\s*\n\s*", " ", cp_para.strip())
            p = doc.add_paragraph(para_text)

            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in p.runs:
                run.font.size = Pt(font_size - 2)  # Assuming font_size is defined

    # --- Section Break for Rest of Front Matter (Starts Section 2) ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    section2 = doc.sections[2]
    # Copy page setup from section 0 to section 2
    section2.page_height = section0.page_height
    section2.page_width = section0.page_width
    section2.left_margin = section0.left_margin  # Outside
    section2.right_margin = section0.right_margin  # Inside
    section2.top_margin = section0.top_margin
    section2.bottom_margin = section0.bottom_margin
    section2.gutter = section0.gutter
    # Ensure mirrored margins are also copied
    sectPr2 = section2._sectPr
    pgMar2 = sectPr2.find(qn("w:pgMar"))
    if pgMar2 is not None:
        pgMar2.set(qn("w:mirrorMargins"), "true")

    # --- Section 2: Rest of Front Matter ---
    logging.info("Adding Rest of Front Matter (Section 2, starts page iii)...")
    fm_order = ["dedication", "foreword", "preface", "acknowledgements"]
    has_fm_content = False
    for key in fm_order:
        content = front_matter.get(key)
        # Check if content exists and is not the placeholder failure message
        if content and not content.startswith(
            f"[{key.title()} content generation failed.]"
        ):
            if has_fm_content:  # Add page break before subsequent FM sections
                doc.add_page_break()
            title = key.replace("_", " ").title()
            # Add space before title is handled by Heading 1 style now
            doc.add_paragraph(title, style="Heading 1")
            # Pass the main doc object as the container for markdown conversion
            context_label = f"FrontMatter_{key}"
            markdown_to_docx(
                content,
                doc,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
                context_label=context_label,
            )
            has_fm_content = True

    # Set page numbering for this section (starts after title/copyright)
    # Roman numerals starting from iii (page 3 conceptually)
    set_page_numbering(section2, format_code="lowerRoman", start_number=3)
    logging.info("Set Front Matter page numbering (lowerRoman, starting iii).")

    # --- Section Break for Body Matter (Starts Section 3) ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    section3 = doc.sections[3]
    # Copy page setup from section 0 to section 3
    section3.page_height = section0.page_height
    section3.page_width = section0.page_width
    section3.left_margin = section0.left_margin  # Outside
    section3.right_margin = section0.right_margin  # Inside
    section3.top_margin = section0.top_margin
    section3.bottom_margin = section0.bottom_margin
    section3.gutter = section0.gutter
    # Ensure mirrored margins are also copied
    sectPr3 = section3._sectPr
    pgMar3 = sectPr3.find(qn("w:pgMar"))
    if pgMar3 is not None:
        pgMar3.set(qn("w:mirrorMargins"), "true")

    # --- Section 3: Body Matter (Starts page 1) ---
    logging.info("Adding Body Matter (Chapters)...")
    chapter_keys = list(body_matter.keys())
    for i, chapter_title in enumerate(chapter_keys):
        sections_data = body_matter[chapter_title]
        logging.info(f"Adding Chapter {i+1}: {chapter_title}")
        # Add chapter title (spacing handled by style)
        doc.add_paragraph(chapter_title, style="Heading 1")

        for j, section_info in enumerate(sections_data):
            section_content = section_info.get("content", "[Missing Content]")
            section_title = section_info.get("title", f"Section {j+1}")
            logging.debug(f"Adding content for Section {j+1}: '{section_title}'")
            # Add section title (spacing handled by style)
            if section_title:
                doc.add_paragraph(section_title, style="Heading 2")

            # --- Construct context label for body section ---
            safe_chap_title = sanitize_filename(chapter_title, 30)
            safe_sec_title = sanitize_filename(section_title, 30)
            context_label = f"Chap{i+1}_{safe_chap_title}_Sec{j+1}_{safe_sec_title}"
            # --- End context label construction ---

            # Pass the main doc object as the container
            markdown_to_docx(
                section_content,
                doc,
                doc,
                config,
                usable_width_inches,
                equation_image_dir,
                context_label=context_label,
            )

        # Add page break after chapter, except for the last one
        if i < len(chapter_keys) - 1:
            doc.add_page_break()

    # Set page numbering for the body section (starts at 1)
    set_page_numbering(section3, format_code="decimal", start_number=1)
    logging.info("Set Body Matter page numbering (decimal, starting at 1).")

    # --- Section Break for Back Matter (Starts Section 4) ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    section4 = doc.sections[4]
    # Copy page setup from section 0 to section 4
    section4.page_height = section0.page_height
    section4.page_width = section0.page_width
    section4.left_margin = section0.left_margin  # Outside
    section4.right_margin = section0.right_margin  # Inside
    section4.top_margin = section0.top_margin
    section4.bottom_margin = section0.bottom_margin
    section4.gutter = section0.gutter
    # Ensure mirrored margins are also copied
    sectPr4 = section4._sectPr
    pgMar4 = sectPr4.find(qn("w:pgMar"))
    if pgMar4 is not None:
        pgMar4.set(qn("w:mirrorMargins"), "true")

    # --- Section 4: Back Matter (Continues numbering) ---
    logging.info("Adding Back Matter (Section 4, continuing numbering)...")
    bm_order = ["appendix", "glossary", "bibliography", "about_the_author"]
    has_bm_content = False
    bm_added_count = 0
    valid_bm_keys = [
        k
        for k in bm_order
        if back_matter.get(k)
        and not back_matter[k].startswith(
            f"[{k.replace('_', ' ').title()} content generation failed.]"
        )
    ]

    for i, key in enumerate(valid_bm_keys):
        content = back_matter[key]
        if has_bm_content:  # Add page break before subsequent BM sections
            doc.add_page_break()
        title = key.replace("_", " ").title()
        # Add title (spacing handled by style)
        doc.add_paragraph(title, style="Heading 1")
        # Pass the main doc object as the container
        context_label = f"BackMatter_{key}"
        markdown_to_docx(
            content,
            doc,
            doc,
            config,
            usable_width_inches,
            equation_image_dir,
            context_label=context_label,
        )
        has_bm_content = True
        bm_added_count += 1

    # Set page numbering for this section (continues from body)
    # Pass None for start_number to continue sequence
    set_page_numbering(section4, format_code="decimal", start_number=None)
    logging.info("Set Back Matter page numbering (continuing decimal).")

    # --- Save Document ---
    try:
        doc.save(output_filename)
        logging.info(f"Main DOCX file assembly complete. Saved as '{output_filename}'")
        return output_filename
    except PermissionError:
        logging.error(
            f"PermissionError: Could not save '{output_filename}'. Check if the file is open or permissions are correct."
        )
        return None  # Indicate failure
    except Exception as e:
        logging.error(f"Error saving main DOCX file '{output_filename}': {e}")
        return None  # Indicate failure


# --- Marketing Docx Assembly (using python-docx) ---
def assemble_marketing_docx(
    config,
    back_matter_content,
    blurb_content,
    summary_context,
    main_book_filename_stem,
    equation_image_dir,
    output_dir,
):
    """
    Assembles the separate marketing DOCX file using python-docx.
    """
    if not main_book_filename_stem:
        logging.error("Cannot create marketing docx without main book filename stem.")
        return
    # Construct the full output path using the provided directory
    output_filename = (
        output_dir / f"{main_book_filename_stem}_Marketing.docx"
    )  # Use pathlib's / operator
    logging.info(f"Assembling marketing DOCX file: '{output_filename}'")
    doc = Document()
    style_config = config.get("style_params", {})
    font_name = style_config.get("font_name", "Times New Roman")
    font_size = style_config.get("font_size", 12)
    gen_params = config.get("generation_params", {})

    try:  # Basic style setup
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)

        # Define Heading 1 if not present
        if "Heading 1" not in doc.styles:
            h1_style = doc.styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
            h1_style.base_style = doc.styles["Normal"]
            h1_style.font.name = font_name
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)
        else:  # Ensure font consistency
            doc.styles["Heading 1"].font.name = font_name

        # Define Heading 2 if not present
        if "Heading 2" not in doc.styles:
            h2_style = doc.styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)
            h2_style.base_style = doc.styles["Normal"]
            h2_style.font.name = font_name
            h2_style.font.size = Pt(13)
            h2_style.font.bold = True
            h2_style.paragraph_format.space_before = Pt(10)
            h2_style.paragraph_format.space_after = Pt(4)
        else:  # Ensure font consistency
            doc.styles["Heading 2"].font.name = font_name

        # Define List Bullet if not present
        if "List Bullet" not in doc.styles:
            lb_style = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
            lb_style.base_style = doc.styles["Normal"]
            # Add basic bullet formatting if needed

    except Exception as e:
        logging.warning(f"Could not apply basic styles to marketing doc: {e}")

    # --- Add Book Details Section ---
    doc.add_paragraph("Book Details", style="Heading 1")

    random_topic_seed = gen_params.get("random_topic_seed")
    if random_topic_seed:
        doc.add_paragraph("Random Topic Seed:", style="Heading 2")
        doc.add_paragraph(random_topic_seed)

    api_settings_local = config.get(
        "api_settings", {}
    )  # Use a local var to avoid conflict
    api_provider_local = api_settings_local.get("provider", "gemini")

    model_name_used = ""
    if api_provider_local == "gemini":
        model_name_used = api_settings_local.get("gemini", {}).get("model", "")
    elif api_provider_local == "ollama":
        model_name_used = api_settings_local.get("ollama", {}).get("model", "")
    else:
        model_name_used = "[Unknown API Provider or Model]"
        logging.warning(
            f"Unknown API provider '{api_provider_local}' when trying to get model name for marketing doc."
        )

    doc.add_paragraph("LLM Model Used:", style="Heading 2")
    doc.add_paragraph(model_name_used)

    doc.add_paragraph("Main Topic:", style="Heading 2")
    doc.add_paragraph(gen_params.get("main_topic", "[Not Specified]"))

    writing_tone = gen_params.get("writing_tone", DEFAULT_WRITING_TONE)
    doc.add_paragraph("Writing Tone:", style="Heading 2")
    doc.add_paragraph(writing_tone)

    doc.add_paragraph("Setting:", style="Heading 2")
    doc.add_paragraph(gen_params.get("setting", "[Not Specified]"))

    doc.add_paragraph("Key Concepts:", style="Heading 2")
    key_concepts = gen_params.get("key_concepts", [])
    if key_concepts:
        for concept in key_concepts:
            # Use List Bullet style, ensure it exists
            list_style = (
                doc.styles["List Bullet"]
                if "List Bullet" in doc.styles
                else doc.styles["Normal"]
            )
            doc.add_paragraph(concept, style=list_style)
    else:
        doc.add_paragraph("[None Specified]")

    # --- Add Character List (if generated) ---
    character_list = gen_params.get("character_list")
    if character_list:
        doc.add_paragraph("characters", style="Heading 2")
        list_style = (
            doc.styles["List Bullet"]
            if "List Bullet" in doc.styles
            else doc.styles["Normal"]
        )
        for char in character_list:
            if isinstance(char, dict) and "name" in char and "description" in char:
                # Add name in bold, then description
                p = doc.add_paragraph(style=list_style)
                p.add_run(f"{char['name']}: ").bold = True
                p.add_run(char["description"])
            else:  # Fallback for unexpected format
                doc.add_paragraph(str(char), style=list_style)
    # --- End Character List ---

    doc.add_paragraph("Author Details:", style="Heading 2")
    doc.add_paragraph(f"Name: {gen_params.get('author_name', '[Not Specified]')}")
    doc.add_paragraph(f"Gender: {gen_params.get('author_gender', '[Not Specified]')}")

    doc.add_page_break()

    # --- Add Book Blurb ---
    doc.add_paragraph("Book Blurb", style="Heading 1")
    markdown_to_docx(
        blurb_content or "[Blurb generation failed]",
        doc,
        doc,
        config,
        None,
        equation_image_dir,
        context_label="Marketing_Blurb",
    )
    doc.add_page_break()

    # --- Add Book Summary Section ---
    doc.add_paragraph("Book Summary (from Chapter Summaries)", style="Heading 1")
    cleaned_summary = (summary_context or "").strip()
    if cleaned_summary and cleaned_summary != "[No chapter summaries available]":
        # Remove the initial label if present
        cleaned_summary = re.sub(
            r"^\s*Chapter summaries:\s*", "", cleaned_summary, flags=re.IGNORECASE
        ).strip()
        markdown_to_docx(
            cleaned_summary,
            doc,
            doc,
            config,
            None,
            equation_image_dir,
            context_label="Marketing_Summary",
        )
    else:
        doc.add_paragraph("[No summary points available]")
    doc.add_page_break()

    # --- Add About the Author ---
    doc.add_paragraph("About the Author", style="Heading 1")
    about_author_content = back_matter_content.get(
        "about_the_author", "[About the Author generation failed]"
    )
    markdown_to_docx(
        about_author_content,
        doc,
        doc,
        config,
        None,
        equation_image_dir,
        context_label="Marketing_AboutAuthor",
    )

    # --- Save the document ---
    try:
        doc.save(output_filename)
        logging.info(
            f"Marketing DOCX file assembly complete. Saved as '{output_filename}'"
        )
    except PermissionError:
        logging.error(
            f"PermissionError: Could not save marketing file '{output_filename}'. Check if file is open or permissions are correct."
        )
    except Exception as e:
        logging.error(f"Error saving marketing DOCX file '{output_filename}': {e}")


# --- Main Execution ---
if __name__ == "__main__":
    logging.info("Starting book generation process...")
    start_time = time.time()
    config = load_config()

    api_settings = config.get("api_settings", {})  # Get api_settings once

    # Determine API provider
    api_provider = api_settings.get("provider", "gemini")
    if api_provider == "gemini":
        api_key = setup_environment()  # This function exits if key not found
        configure_gemini(api_key)  # This function exits on error
        logging.info("Gemini API provider configured.")
    elif api_provider == "ollama":
        ollama_settings = config.get("ollama_settings", {})
        ollama_base_url = ollama_settings.get("base_url", "http://localhost:11434")
        try:
            logging.info(
                f"Ollama API provider selected. Attempting to connect to: {ollama_base_url}"
            )
            # Quick health check for Ollama server
            requests.get(
                f"{ollama_base_url.rstrip('/')}/api/tags", timeout=5
            ).raise_for_status()  # List models as a basic check
            logging.info(
                f"Successfully connected to Ollama server at {ollama_base_url}"
            )
        except requests.exceptions.RequestException as e:
            logging.error(
                f"Error: Could not connect or communicate with Ollama server at {ollama_base_url}. Error: {e}. Please ensure Ollama is running and accessible."
            )
            sys.exit(1)
    else:
        logging.error(
            f"Unsupported API provider '{api_provider}' specified in config. Supported: 'gemini', 'ollama'. Exiting."
        )
        sys.exit(1)

    # --- Determine Output Directory ---
    # Default to a subdirectory named 'output' in the current working directory
    output_base_dir_str = config.get("output_dir", "output")
    output_base_dir = pathlib.Path(output_base_dir_str)
    # Create the base output directory if it doesn't exist
    try:
        output_base_dir.mkdir(parents=True, exist_ok=True)
        logging.info(f"Using output directory: {output_base_dir.resolve()}")
    except Exception as e:
        logging.error(
            f"Failed to create output directory '{output_base_dir}': {e}. Exiting."
        )
        sys.exit(1)
    # Store the resolved path back in config for potential use elsewhere (optional)
    config["output_dir_resolved"] = output_base_dir
    # --- End Output Directory Setup ---

    # --- Determine Base Cache Directory ---
    base_cache_dir_from_config = api_settings.get("base_cache_dir", "api_cache")
    pathlib.Path(base_cache_dir_from_config).mkdir(parents=True, exist_ok=True)
    logging.info(f"Using base cache directory: {base_cache_dir_from_config}")

    # --- Determine Main Topic ---
    generation_params = config.setdefault("generation_params", {})  # Ensure exists

    # Check if main_topic is provided in the config
    if not generation_params.get("main_topic"):
        logging.info(
            "No 'main_topic' found in config. Attempting to auto-generate one."
        )
        random_topic = generate_random_topic(config)  # API call happens here
        if random_topic:
            generation_params["main_topic"] = random_topic
            logging.info(f"Auto-generated main_topic: '{random_topic}'")
        else:
            # If generation fails and it wasn't in config, we cannot proceed.
            logging.critical(
                "Fatal: Failed to auto-generate random topic and none provided in config. Exiting."
            )
            sys.exit(1)
    else:
        # If main_topic was provided in config, use it.
        logging.info(
            f"Using main_topic from config: '{generation_params['main_topic']}'"
        )
    # --- End Main Topic Determination ---

    # --- Construct Topic-Specific Cache Path AFTER topic is determined ---
    main_topic = generation_params["main_topic"]
    sanitized_topic = sanitize_filename(main_topic, 64)
    # Generate a short hash of the original topic for uniqueness
    topic_hash = hashlib.sha1(main_topic.encode("utf-8")).hexdigest()[:8]
    topic_dir_name = f"{sanitized_topic}_{topic_hash}"

    # --- Get and sanitize the model name for cache path ---
    if api_provider == "gemini":
        gemini_conf = api_settings.get("gemini", {})
        model_name_for_cache = gemini_conf.get("model", "gemini-2.0-flash-latest")
    elif api_provider == "ollama":
        ollama_conf = api_settings.get("ollama", {})
        model_name_for_cache = ollama_conf.get("model", "ollama_default_model")
    else:  # Should not happen due to earlier check
        model_name_for_cache = "unknown_api_provider_model"
    # Sanitize the model name to make it directory-safe, limit length
    sanitized_model_name = sanitize_filename(model_name_for_cache, 30)
    logging.info(f"Using model name for cache path: '{sanitized_model_name}'")

    # --- Create the cache path including the model name ---
    topic_specific_cache_dir = (
        pathlib.Path(base_cache_dir_from_config) / sanitized_model_name / topic_dir_name
    )  # Use pathlib
    logging.info(f"Topic-specific cache directory set to: {topic_specific_cache_dir}")
    # Update the config dictionary IN MEMORY so subsequent calls to call_llm_api use the right path
    config["cache_dir"] = str(
        topic_specific_cache_dir
    )  # Store as string if needed elsewhere
    # Ensure the specific directory exists
    topic_specific_cache_dir.mkdir(parents=True, exist_ok=True)

    # --- Define and Create Equation Image Directory HERE ---
    # Place equation images within the topic-specific, model-specific cache dir
    equation_image_dir = topic_specific_cache_dir / "equation_images"
    equation_image_dir.mkdir(parents=True, exist_ok=True)
    logging.info(f"Equation image directory set to: {equation_image_dir}")
    # --- End Equation Image Directory Setup ---

    # Determine Setting
    if not generation_params.get("setting"):
        logging.info("No 'setting' found in config. Attempting to auto-generate one.")
        generated_setting = generate_setting(config)  # API call
        if generated_setting:
            generation_params["setting"] = generated_setting
            logging.info(f"Auto-generated setting: '{generated_setting[:100]}...'")
        else:
            # If generation fails and it wasn't in config, use a placeholder.
            logging.warning("Failed to auto-generate setting. Using placeholder.")
            generation_params["setting"] = "[Setting Generation Failed]"
    else:
        # If setting was provided in config, use it.
        logging.info(
            f"Using setting from config: '{generation_params['setting'][:100]}...'"
        )

    # Ensure setting exists even if generation failed or config was empty initially
    generation_params.setdefault(
        "setting", "[No Setting Provided or Generation Failed]"
    )
    # Log the final setting being used (could be from config, generated, or placeholder)
    logging.info(f"Final setting being used: '{generation_params['setting'][:100]}...'")

    # Determine Key Concepts: Use from config if provided, otherwise generate.
    key_concepts = generation_params.get("key_concepts")

    # Validate if key_concepts from config is a list and non-empty
    if key_concepts and isinstance(key_concepts, list):
        logging.info(
            f"Using {len(key_concepts)} key concepts provided in configuration."
        )
        # Ensure the list contains only strings and strip whitespace
        generation_params["key_concepts"] = [
            str(item).strip() for item in key_concepts if str(item).strip()
        ]
        # Re-check length after cleaning
        if not generation_params["key_concepts"]:
            logging.warning(
                "Provided key_concepts list was empty after cleaning. Will attempt to generate."
            )
            key_concepts = None  # Trigger generation below
        else:
            logging.info(
                f"Validated {len(generation_params['key_concepts'])} key concepts from configuration."
            )

    else:
        if key_concepts is not None:  # Log if it was present but invalid
            logging.warning(
                f"Invalid or empty 'key_concepts' found in config (type: {type(key_concepts)}). Attempting auto-generation."
            )
        else:  # Log if it was missing entirely
            logging.info(
                "No 'key_concepts' found in config. Attempting auto-generation."
            )
        key_concepts = None  # Ensure generation is triggered

    # Attempt generation if needed
    if key_concepts is None:
        logging.info("Attempting to auto-generate key concepts.")
        generated_concepts = generate_key_concepts(config)  # API call
        if generated_concepts:
            generation_params["key_concepts"] = generated_concepts
            # Format the concepts with newlines for logging
            concepts_formatted_for_log = "\n".join(
                f"- {concept}" for concept in generated_concepts
            )
            # Log the count and the formatted list
            logging.info(
                f"Successfully auto-generated {len(generated_concepts)} key concepts:\n{concepts_formatted_for_log}"
            )
        else:
            logging.warning(
                "Failed to auto-generate key concepts. Proceeding with an empty list."
            )
            generation_params[
                "key_concepts"
            ] = []  # Ensure it's an empty list on failure

    # Final log of the count being used
    final_count = len(generation_params.get("key_concepts", []))
    logging.info(f"Using {final_count} key concepts for generation.")

    # --- Determine Author Name and Gender ---
    author_name = generation_params.get("author_name", "").strip()
    author_gender = generation_params.get("author_gender", "").strip().lower()
    valid_genders = ["male", "female", "other"]

    if not author_name and not author_gender:
        logging.info("Author name and gender missing. Generating both...")
        # 1. Generate gender first
        generated_gender = generate_random_gender(config)  # Simple random choice
        logging.info(f"Randomly selected gender: {generated_gender}")
        # 2. Generate name based on gender
        generated_name = generate_random_name(config, generated_gender)  # API call
        if generated_name:
            author_name = generated_name
            author_gender = generated_gender
            generation_params["author_name"] = author_name
            generation_params["author_gender"] = author_gender
            logging.info(
                f"Generated Author: Name='{author_name}', Gender='{author_gender}'"
            )
        else:
            logging.critical(
                "Fatal: Failed to generate author name when both name and gender were missing. Exiting."
            )
            sys.exit(1)

    elif not author_name:
        logging.info(
            f"Author name missing. Generating name for specified gender: '{author_gender}'..."
        )
        if author_gender not in valid_genders:
            logging.warning(
                f"Provided gender '{author_gender}' is not standard ({valid_genders}). Attempting name generation anyway."
            )
            # Decide if you want to default the gender here or proceed. Let's proceed.

        generated_name = generate_random_name(config, author_gender)  # API call
        if generated_name:
            author_name = generated_name
            generation_params["author_name"] = author_name
            logging.info(
                f"Generated Author Name: '{author_name}' (Gender was '{author_gender}')"
            )
        else:
            logging.critical(
                f"Fatal: Failed to generate author name for gender '{author_gender}'. Exiting."
            )
            sys.exit(1)

    elif not author_gender:
        logging.info(
            f"Author gender missing. Attempting to determine gender from name: '{author_name}'..."
        )
        determined_gender = determine_gender_from_name(config, author_name)  # API call
        if determined_gender:
            author_gender = determined_gender
            generation_params["author_gender"] = author_gender
            logging.info(
                f"Determined Author Gender: '{author_gender}' (Name was '{author_name}')"
            )
        else:
            # If determination fails, you need a fallback. Using 'other' or exiting are options.
            logging.warning(
                f"Could not determine gender for '{author_name}'. Falling back to 'other'."
            )
            author_gender = "other"
            generation_params["author_gender"] = author_gender
            # Or, make it critical:
            # logging.critical(f"Fatal: Could not determine gender for name '{author_name}'. Exiting.")
            # sys.exit(1)

    else:
        # Both name and gender were provided
        logging.info(f"Using Author Name from config: '{author_name}'")
        if author_gender not in valid_genders:
            logging.warning(
                f"Author gender ('{author_gender}') from config is not standard ({valid_genders}). Using it anyway."
            )
        logging.info(f"Using Author Gender from config: '{author_gender}'")

    # Final check - ensure both have values before proceeding (should always pass if logic above is correct)
    if not generation_params.get("author_name"):
        logging.critical(
            "Fatal: Author name is still missing after processing. Exiting."
        )
        sys.exit(1)
    if not generation_params.get("author_gender"):
        logging.critical(
            "Fatal: Author gender is still missing after processing. Exiting."
        )
        sys.exit(1)

    logging.info(
        f"Final Author Details: Name='{generation_params['author_name']}', Gender='{generation_params['author_gender']}'"
    )
    # --- End Author Name and Gender Determination ---

    # Determine Writing Tone
    # Check if writing_tone is provided and not empty in the config
    writing_tone = generation_params.get("writing_tone", "").strip()

    if not writing_tone:
        # If not provided or empty, attempt to auto-generate it
        logging.info(
            "No 'writing_tone' found or it was empty in config. Attempting to auto-generate one."
        )
        generated_tone = generate_writing_tone(config)  # API call
        if generated_tone:
            writing_tone = generated_tone
            generation_params[
                "writing_tone"
            ] = generated_tone  # Update config in memory
            logging.info(f"Auto-generated writing tone: '{writing_tone}'")
        else:
            # If generation fails, fall back to the default
            writing_tone = DEFAULT_WRITING_TONE
            generation_params[
                "writing_tone"
            ] = writing_tone  # Store default back in config
            logging.warning(
                f"Failed to auto-generate writing tone. Using default: '{writing_tone}'"
            )
    else:
        # If it was provided in the config, use that value
        logging.info(f"Using writing_tone from config: '{writing_tone}'")

    # Ensure writing_tone has a value (either from config, generated, or default)
    if (
        not writing_tone
    ):  # Should ideally not happen due to fallback, but as a safeguard
        writing_tone = DEFAULT_WRITING_TONE
        generation_params["writing_tone"] = writing_tone
        logging.warning(
            f"Writing tone was still empty after checks. Using default: '{writing_tone}'"
        )

    logging.info(f"Final writing tone being used: '{writing_tone}'")
    # --- End Writing Tone Determination ---

    # --- Generate Core Book Structure ---
    # Check if book_title is provided in the config
    book_title_from_config = generation_params.get("book_title", "").strip()

    if book_title_from_config:
        book_title = book_title_from_config
        logging.info(f"Using Book Title from config: '{book_title}'")
    else:
        logging.info(
            "No 'book_title' found in config or it was empty. Attempting to auto-generate one."
        )
        book_title = generate_book_title(config)  # API call, exits on failure
        # generate_book_title exits if it fails or gets an empty title.
        # If it returns, it's a valid, cleaned title.
        logging.info(f"Successfully auto-generated Book Title: {book_title}")
        # Optionally, store the generated title back into config if desired for consistency
        # This ensures that if other parts of the code were to re-read it from config,
        # they'd see the generated one.
        generation_params["book_title"] = book_title

    # --- Determine if the book is fiction ---
    # This call will now use internal caching and set generation_params["is_fiction"]
    is_book_non_fiction(
        config, book_title
    )  # Ensures generation_params["is_fiction"] is set

    if generation_params.get(
        "is_fiction", False
    ):  # Default to False (non-fiction) if key is somehow missing
        logging.info(
            f"Book '{book_title}' has been identified as fiction. Detailed section generation within chapters will be skipped; chapter content will be generated as a single block."
        )
    else:
        logging.info(
            f"Book '{book_title}' has been identified as non-fiction (or type indeterminate, proceeding with detailed sections). Detailed sections will be generated within chapters."
        )
    # --- End Fiction Determination ---

    # --- Generate Character List (if enabled) ---
    # This needs book_title, topic, setting, concepts
    generate_character_list(config, book_title)  # API call inside if enabled
    # The result (or None) is stored in config['generation_params']['character_list']

    # --- Generate Chapter Outline (potentially using characters) ---
    # Prepare character context string
    character_context_for_prompts = format_character_list_for_prompt(
        config["generation_params"].get("character_list")
    )

    chapter_titles = generate_chapter_outline(config, character_context_for_prompts)

    if chapter_titles:
        formatted_outline = "\n".join(
            f"{i+1}. {title}" for i, title in enumerate(chapter_titles)
        )
        logging.info(f"Generated Chapter Outline:\n{formatted_outline}")
    else:
        logging.error("Chapter outline generation resulted in no titles. Exiting.")
        sys.exit(1)  # Cannot proceed without chapters

    # --- Pass 1: Generate Chapter Summaries ---
    logging.info("--- Starting Pass 1: Generating Chapter Summaries ---")
    chapter_summaries = {}
    previous_summaries_list = []  # Keep track of summaries generated so far

    for i, chap_title in enumerate(chapter_titles):
        logging.info(f"--- Generating Summary for Chapter {i+1}: {chap_title} ---")

        # --- Create context string from previous summaries ---
        previous_summaries_context = ""
        if previous_summaries_list:
            # Format the context clearly, e.g., as a list
            context_parts = [
                f"- Chapter {idx+1}: {s}"
                for idx, s in enumerate(previous_summaries_list)
            ]
            previous_summaries_context = "\n".join(context_parts)
            logging.debug(
                f"Providing context of {len(previous_summaries_list)} previous summaries for chapter '{chap_title}'."
            )
        # --- End context creation ---

        # Pass the context to the generation function
        summary = generate_chapter_summary(
            config,
            chap_title,
            writing_tone,
            previous_summaries_context,
            character_context_for_prompts,
        )
        chapter_summaries[chap_title] = summary

        # --- Add the generated summary to the list for the next iteration's context ---
        # Only add valid, non-placeholder summaries to the context for subsequent chapters
        cleaned_summary_for_context = summary.strip()
        if cleaned_summary_for_context and not cleaned_summary_for_context.startswith(
            "Placeholder summary"
        ):
            previous_summaries_list.append(cleaned_summary_for_context)
        # --- End adding summary to list ---

    # --- Prepare Summary Context for subsequent prompts ---
    summary_context = "[No chapter summaries available]"  # Default
    summary_parts = []
    for i, chap_title in enumerate(chapter_titles):
        summary = chapter_summaries.get(chap_title, "").strip()
        if summary and not summary.startswith("Placeholder summary"):
            summary_parts.append(f"Chapter {i+1} ('{chap_title}'): {summary}")
    if summary_parts:
        summary_context = "Chapter summaries:\n" + "\n\n".join(summary_parts)
        logging.debug(f"Generated summary context:\n{summary_context}")
    else:
        logging.warning(
            "No valid chapter summaries generated to create summary context."
        )

    # --- Generate Overall Summary and Save to Markdown ---
    overall_summary_text = generate_overall_summary(config, book_title, summary_context)
    save_summary_to_markdown(book_title, overall_summary_text, output_base_dir)
    # --- End Overall Summary Generation ---

    # --- Generate Front/Back Matter & Marketing Content (using summaries) ---
    logging.info("--- Generating Front Matter, Back Matter, and Marketing Content ---")
    front_matter_content = generate_front_matter(
        config,
        book_title,
        generation_params["author_name"],
        writing_tone,
        summary_context,
    )  # Multiple API calls inside

    back_matter_content = generate_back_matter(
        config,
        book_title,
        generation_params["author_name"],
        generation_params["author_gender"],
        writing_tone,
        summary_context,
    )  # Multiple API calls inside

    main_filename_stem = sanitize_filename(book_title)
    if main_filename_stem != "sanitized_empty":
        blurb_text = generate_book_blurb(
            config, book_title, summary_context, writing_tone
        )  # API call
        assemble_marketing_docx(
            config,
            back_matter_content,
            blurb_text,
            summary_context,
            main_filename_stem,
            equation_image_dir,
            output_base_dir,
        )
    else:
        logging.error(
            f"Could not generate valid filename stem from title '{book_title}'. Skipping marketing DOCX."
        )

    # --- Pass 2: Generate Section Content (using summaries) ---
    logging.info("--- Starting Pass 2: Generating Section Content ---")
    body_matter = {}
    is_fiction_book = generation_params.get("is_fiction", False)  # Get the flag

    for i, chap_title in enumerate(chapter_titles):
        logging.info(f"--- Processing Chapter {i+1}: {chap_title} ---")
        chapter_summary = chapter_summaries.get(
            chap_title, f"Placeholder summary for chapter '{chap_title}'."
        )  # Use placeholder if missing
        body_matter[chap_title] = []  # Initialize with empty list for sections

        if is_fiction_book:
            logging.info(
                f"Book is fiction. Generating content for chapter '{chap_title}' as a single block."
            )
            # Use a generic title for the single content block of a fiction chapter
            fiction_section_title = "Narrative"  # Or "Chapter Content"
            chapter_content_as_single_block = generate_section_content(
                config,
                chap_title,  # chapter_title
                fiction_section_title,  # section_title (generic for fiction)
                1,  # section_num
                1,  # total_sections
                chapter_summary,
                writing_tone,
                character_context_for_prompts,
            )
            body_matter[chap_title].append(
                {
                    "title": "",
                    "content": chapter_content_as_single_block,
                }
            )
            if (
                not chapter_content_as_single_block
                or "Content generation failed" in chapter_content_as_single_block
            ):
                logging.warning(
                    f"Content generation potentially failed for fiction chapter '{chap_title}' (treated as single section)."
                )
        else:  # Non-fiction book, proceed with normal section generation
            # --- Call generate_section_titles with added context ---
            section_titles = generate_section_titles(
                config,
                chap_title,
                chapter_summary,
                chapter_titles,
                chapter_summaries,
                character_context_for_prompts,
            )

            if not section_titles:
                logging.warning(
                    f"No section titles generated for chapter '{chap_title}'. Adding chapter summary as placeholder content."
                )
                # Add chapter summary as content if no sections are generated
                body_matter[chap_title].append(
                    {"title": "Chapter Overview", "content": chapter_summary}
                )
                continue  # Skip to next chapter
            else:
                formatted_sections = "\n".join(
                    f"  {j+1}. {title}" for j, title in enumerate(section_titles)
                )
                logging.info(
                    f"Generated {len(section_titles)} section titles for chapter '{chap_title}':\n{formatted_sections}"
                )

            for j, sec_title in enumerate(section_titles):
                section_content = generate_section_content(
                    config,
                    chap_title,
                    sec_title,
                    j + 1,
                    len(section_titles),
                    chapter_summary,
                    writing_tone,
                    character_context_for_prompts,
                )
                body_matter[chap_title].append(
                    {"title": sec_title, "content": section_content}
                )
                if (
                    not section_content
                    or "Content generation failed" in section_content
                ):
                    logging.warning(
                        f"Content generation potentially failed for Chapter '{chap_title}', Section '{sec_title}'."
                    )

    logging.info("Finished generating all body matter.")

    # --- Assemble Final DOCX ---
    main_docx_filename = assemble_docx(
        config,
        front_matter_content,
        body_matter,
        back_matter_content,
        book_title,
        equation_image_dir,
        output_base_dir,
    )

    end_time = time.time()
    logging.info(
        f"Book generation process finished in {end_time - start_time:.2f} seconds."
    )
    if main_docx_filename:
        logging.info(f"Main book saved as: {main_docx_filename}")
    else:
        logging.error("Main book DOCX file failed to save.")
